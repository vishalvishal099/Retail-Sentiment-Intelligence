"""
Build the BITS WILP Mid-Semester Progress Report (.docx) for:
    Vishal Singh, 2020AA05641
    Real-Time Social Media Mining and Trust-Aware Sentiment Analysis
    Using Large Language Models for Retail Product Feedback Optimization

Layout mirrors the supplied sample (`MID SEM REPORT - Sample copy - 29.07.2021 (1).pdf`):
title pages -> abstract -> contents -> body sections -> future plan -> abbreviations.

Output:
    docs/Sem_4/MID_SEM_REPORT_VishalSingh_2020AA05641.docx
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_PATH = Path(__file__).parent / "MID_SEM_REPORT_VishalSingh_2020AA05641.docx"
FIGURES_DIR = Path(__file__).parent / "figures"


# ---------- helpers ----------

def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=110, right=110) -> None:
    """Set cell padding in twips (1/20 pt). Defaults give ~4pt vertical, ~5.5pt horizontal."""
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def _set_outline_level(paragraph, level: int) -> None:
    """Mark paragraph with an outline level so Word's Navigation Pane / TOC sees it."""
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


_BOOKMARK_ID = [1000]


def _add_bookmark(paragraph, name: str) -> None:
    """Wrap paragraph contents with a Word bookmark so hyperlinks can target it."""
    bm_id = str(_BOOKMARK_ID[0])
    _BOOKMARK_ID[0] += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bm_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bm_id)
    # bookmarkStart goes after pPr (if present) but before any runs
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is not None:
        p_pr.addnext(start)
    else:
        paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_toc_link(doc: Document, text: str, bookmark: str,
                  indent: bool = False, size: int = 11) -> None:
    """Add a clickable TOC entry that jumps to the given bookmark when clicked."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.25
    if indent:
        pf.left_indent = Inches(0.35)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFont = OxmlElement("w:rFonts")
    rFont.set(qn("w:ascii"), "Calibri")
    rFont.set(qn("w:hAnsi"), "Calibri")
    rpr.append(rFont)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2))
    rpr.append(sz)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
    rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hyperlink.append(r)
    p._p.append(hyperlink)


def add_caption(doc: Document, text: str, bookmark: str | None = None) -> None:
    """Bold caption above a table (e.g. 'Table 1: ...'); optionally bookmarked."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.2
    pf.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    if bookmark:
        _add_bookmark(p, bookmark)


def add_heading(doc: Document, text: str, level: int = 1,
                bookmark: str | None = None) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 10)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    if level == 1:
        run.font.size = Pt(15)
    elif level == 2:
        run.font.size = Pt(12.5)
    else:
        run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x04, 0x1E, 0x42)  # walmart navy
    _set_outline_level(p, max(0, level - 1))
    if bookmark:
        _add_bookmark(p, bookmark)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
             size: int = 11, align=None, space_after: int = 6,
             justify: bool = True) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.2
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    pf = p.paragraph_format
    pf.space_after = Pt(3)
    pf.line_spacing = 1.2
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None, header_bg: str = "041E42",
              zebra: str = "F2F4F7") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr_cells[i], header_bg)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=110, right=110)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # body
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        zebra_on = (r_idx % 2 == 0)
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.18
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cells[c_idx], top=70, bottom=70, left=110, right=110)
            if zebra_on:
                set_cell_bg(cells[c_idx], zebra)

    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(w)


def add_image(doc: Document, filename: str, caption: str, width_in: float = 6.3) -> None:
    """Insert a centered figure image followed by a centered italic caption."""
    path = FIGURES_DIR / filename
    if not path.exists():
        add_para(doc,
                 f"[Missing figure: {filename} — run docs/Sem_4/generate_diagrams.py]",
                 italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(2)
    pf.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.line_spacing = 1.15
    crun = cap.add_run(caption)
    crun.italic = True
    crun.font.name = "Calibri"
    crun.font.size = Pt(10)
    crun.font.color.rgb = RGBColor(0x04, 0x1E, 0x42)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def _add_page_number_field(paragraph) -> None:
    """Insert a 'Page X of Y' field pair into the given paragraph."""
    def _fld(instr: str):
        r = OxmlElement("w:r")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        # placeholder run inside the field so Word renders a default value
        inner = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18")
        rpr.append(sz)
        inner.append(rpr)
        txt = OxmlElement("w:t")
        txt.text = "1"
        inner.append(txt)
        fld.append(inner)
        return fld

    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x5A, 0x64, 0x70)
    paragraph._p.append(_fld("PAGE"))
    run2 = paragraph.add_run(" of ")
    run2.font.size = Pt(9)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(0x5A, 0x64, 0x70)
    paragraph._p.append(_fld("NUMPAGES"))


def setup_page_chrome(doc: Document) -> None:
    """No running header. Footer contains only a centered page number."""
    for section in doc.sections:
        section.top_margin = Inches(0.95)
        section.bottom_margin = Inches(0.95)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.45)

        # ----- header: empty -----
        hp = section.header.paragraphs[0]
        hp.text = ""

        # ----- footer: centered page number only -----
        fp = section.footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        inner = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20")
        rpr.append(sz)
        inner.append(rpr)
        txt = OxmlElement("w:t")
        txt.text = "1"
        inner.append(txt)
        fld.append(inner)
        fp._p.append(fld)


# ---------- build ----------

def build() -> None:
    doc = Document()

    # base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.25
    pf.space_after = Pt(6)

    # page margins
    for section in doc.sections:
        section.top_margin = Inches(0.95)
        section.bottom_margin = Inches(0.95)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # running header + centered footer page numbers on every page
    setup_page_chrome(doc)

    # ===== TITLE PAGE 1 =====
    add_para(doc, "", space_after=24)
    add_para(doc,
             "REAL-TIME SOCIAL MEDIA MINING AND TRUST-AWARE\n"
             "SENTIMENT ANALYSIS USING LARGE LANGUAGE MODELS\n"
             "FOR RETAIL PRODUCT FEEDBACK OPTIMIZATION",
             bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "(Beyond the Stars)", italic=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "BITS ZG628T: Dissertation — Mid-Semester Progress Report",
             bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc, "by", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, "Vishal Singh", bold=True, size=14,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "ID No. 2020AA05641", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "Dissertation work carried out at",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Walmart Global Tech, Bengaluru, India", bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "Submitted in partial fulfilment of",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "M.Tech in Artificial Intelligence & Machine Learning",
             bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(doc, "Under the Supervision of",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Mr. Varunendra Pratap Singh", bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Principal Software Engineer, Walmart Global Tech, Bengaluru",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE",
             bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "PILANI (RAJASTHAN)", bold=True, size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "June 2026", size=12,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_page_break(doc)

    # ===== ABSTRACT =====
    add_heading(doc, "ABSTRACT", level=1, bookmark="sec_abstract")
    add_para(doc,
        "Public social media has become a primary, real-time channel for retail customer voice. "
        "DataReportal (Jan 2024) reports about 5.04 billion active social users, Reddit's S-1 "
        "filing (Feb 2024) discloses ~267 million weekly active uniques, and Sprout Social's 2024 "
        "Index finds 51% of consumers expect a same-day brand response on social. At the same "
        "time, the World Economic Forum (2024) estimates that fake online reviews influence "
        "approximately USD 152 billion in purchases each year, with 4%–15% of reviews flagged as "
        "fake across studies. Any analytical layer that consumes this signal must therefore filter "
        "noise and low-credibility content explicitly.")
    add_para(doc,
        "This dissertation builds a working, end-to-end prototype that ingests public Reddit "
        "posts from a curated set of 32 retail-relevant communities, applies sentiment "
        "classification with a fine-tuned ModernBERT model and aspect-based opinion mining "
        "with a zero-shot DeBERTa-v3 NLI classifier over a fixed retail taxonomy (pricing, "
        "product quality, customer service, store experience, online/app, delivery/pickup, "
        "returns, account/login), captions image attachments with Gemma 3 4B (Ollama) using a "
        "multi-pass prompt pipeline informed by recent vision-language literature, and assigns "
        "each post a 0–1 trust score derived from account-metadata heuristics combined with a "
        "rule-based credibility scorer. Trust-filtered, aspect-tagged results are aggregated "
        "and presented in a React-based dashboard with a human-in-the-loop review queue and a "
        "priority-negatives panel (P1 / P2) that ranks trustworthy negative posts by "
        "trust × confidence.")
    add_para(doc,
        "This mid-semester report documents the work completed between 25 April 2026 and "
        "21 June 2026, covering the literature review, system design, data ingestion and "
        "pre-processing module, and the model analysis (sentiment, aspect, vision) and "
        "trust-scoring module. The remaining "
        "phases — aggregation and reporting refinements, evaluation on a manually labelled sample "
        "of ~250–300 posts, dissertation review, and final submission — are planned for the "
        "second half of the semester. The prototype uses only public data and is not deployed "
        "inside Walmart systems.")
    add_para(doc, "", space_after=18)
    # signature block
    sig = doc.add_table(rows=2, cols=2)
    sig.autofit = True
    sig.cell(0, 0).text = "Signature of the Student"
    sig.cell(0, 1).text = "Signature of the Supervisor"
    sig.cell(1, 0).text = "Name: Vishal Singh\nDate:\nPlace: Bengaluru"
    sig.cell(1, 1).text = "Name: Varunendra Pratap Singh\nDate:\nPlace: Bengaluru"
    for row in sig.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    add_page_break(doc)

    # ===== CONTENTS =====
    add_heading(doc, "Contents", level=1, bookmark="sec_contents")
    toc_entries = [
        ("1. Introduction and Broad Area of Work", "sec_1", False),
        ("2. Problem Statement and Objectives", "sec_2", False),
        ("3. System Design and Architecture", "sec_3", False),
        ("4. Implementation Progress (25 April 2026 – 21 June 2026)", "sec_4", False),
        ("4.1  Data Ingestion and Pre-processing", "sec_4_1", True),
        ("4.2  Model Analysis Module (Sentiment, Aspects, Vision)", "sec_4_2", True),
        ("4.3  Trust-Score Module", "sec_4_3", True),
        ("4.4  Storage Layer", "sec_4_4", True),
        ("4.5  Dashboard and Human-in-the-Loop Review", "sec_4_5", True),
        ("5. Tools, Technologies, and Configuration", "sec_5", False),
        ("6. Problems Encountered and Mitigations", "sec_6", False),
        ("7. Preliminary Observations", "sec_7", False),
        ("8. Future Plan (Remaining Phases)", "sec_8", False),
        ("9. Abbreviations", "sec_9", False),
        ("10. References", "sec_10", False),
    ]
    for text, bm, indent in toc_entries:
        add_toc_link(doc, text, bm, indent=indent)
    add_para(doc, "", space_after=8)
    add_para(doc, "List of Figures and Tables", bold=True, size=11, space_after=4)
    figs_tables = [
        ("Figure 1: Layered System Architecture", "fig_1"),
        ("Figure 2: End-to-End Pipeline Flow (single tick)", "fig_2"),
        ("Figure 3: Trust-Score Composition and P1/P2 Tier Rules", "fig_3"),
        ("Figure 4: Dashboard Information Architecture and Feedback Loop", "fig_4"),
        ("Table 1: Curated Subreddit Coverage (snapshot)", "tbl_1"),
        ("Table 2: Eight-Aspect Retail Taxonomy", "tbl_2"),
        ("Table 3: Tools and Technologies Used", "tbl_3"),
        ("Table 4: Problems Encountered and Mitigations", "tbl_4"),
        ("Table 5: Future Plan with Status", "tbl_5"),
    ]
    for text, bm in figs_tables:
        add_toc_link(doc, text, bm)
    add_page_break(doc)

    # ===== 1. INTRODUCTION =====
    add_heading(doc, "1. Introduction and Broad Area of Work", level=1, bookmark="sec_1")
    add_para(doc,
        "The broad area of this dissertation is Applied Natural Language Processing (NLP) for "
        "analysis of public retail customer feedback on social media. The work spans four "
        "sub-areas: (a) supervised text classification with transformer encoders for sentiment "
        "in short, informal English social-media text; (b) zero-shot natural-language-inference "
        "models for aspect-based opinion mining over a fixed retail taxonomy; (c) multimodal "
        "vision-language models for captioning and OCR of screenshots and photos attached to "
        "retail complaints; and (d) lightweight rule-based trust / credibility filtering using "
        "account-metadata heuristics and retail-insider terminology signals, with data "
        "engineering for scheduled ingestion of public posts and structured storage suitable for "
        "aggregation and reporting.")
    add_para(doc,
        "The motivation is volume and noise. Customers regularly post about their Walmart, "
        "Sam's Club, Spark-driver, and OGP/pickup experiences on Reddit. In the candidate's "
        "current role at Walmart Global Tech (order-fulfillment / product platform), this signal "
        "is read via ad-hoc browsing and keyword-filter dashboards, which suffer from poor "
        "coverage, noise, and latency. The prototype built in this dissertation demonstrates that "
        "a stack of task-specific models — a fine-tuned encoder for sentiment, a zero-shot NLI "
        "model for aspects, and a multimodal vision model for images — combined with an "
        "explicit trust filter can convert this raw public stream into a structured, "
        "aspect-tagged, trust-weighted feed suitable for downstream operational review.")

    # ===== 2. PROBLEM STATEMENT =====
    add_heading(doc, "2. Problem Statement and Objectives", level=1, bookmark="sec_2")
    add_para(doc,
        "Manual monitoring of retail-related Reddit communities cannot keep pace with the daily "
        "volume of posts, and conventional rule-based or lexicon-based sentiment systems perform "
        "poorly on short, sarcastic, slang-heavy retail text. Furthermore, fake or low-credibility "
        "posts contaminate any aggregate signal if consumed without filtering.")
    add_para(doc, "Objectives for the dissertation (as approved in the abstract):", bold=True)
    add_numbered(doc, "Survey recent literature on LLM-based sentiment and aspect-based opinion "
                      "mining and on fake-review / bot-credibility detection.")
    add_numbered(doc, "Build a data ingestion pipeline that periodically collects retail-related "
                      "posts from Reddit via the official Reddit API and a public historical archive.")
    add_numbered(doc, "Implement a model-based analysis module that assigns each post a "
                      "sentiment label and one or more aspects from a fixed retail taxonomy, plus "
                      "a vision caption when an image is attached.")
    add_numbered(doc, "Implement a trust score (0–1) per post combining account-metadata "
                      "heuristics with a rule-based credibility scorer (with an optional "
                      "cloud-LLM credibility path for cost-sensitive deployments).")
    add_numbered(doc, "Produce structured aggregated outputs (top aspects, sentiment distribution "
                      "per aspect, representative examples) with low-trust posts excluded.")
    add_numbered(doc, "Evaluate the prototype on a manually labelled sample of ~250–300 posts: "
                      "sentiment macro-F1, aspect macro-F1, and ROC-AUC for the trust filter.")
    add_numbered(doc, "Document design, evaluation, limitations, and recommended next steps in "
                      "the final dissertation.")

    # ===== 3. SYSTEM DESIGN =====
    add_heading(doc, "3. System Design and Architecture", level=1, bookmark="sec_3")
    add_para(doc,
        "The system is built as a modular pipeline so that each stage can be developed, "
        "evaluated, and replaced independently. Figure 1 shows the five logical layers and "
        "the components inside each, and Figure 2 traces a single end-to-end tick through "
        "those layers.")

    add_heading(doc, "Figure 1: Layered System Architecture", level=3, bookmark="fig_1")
    add_image(doc, "fig1_architecture.png",
              "Figure 1: Five-layer architecture — data sources, ingestion, trust + LLM "
              "analysis, storage / aggregation, and the dashboard surface.")
    add_para(doc,
        "In words: Reddit Sources → Ingestion Layer (PRAW live + Arctic-Shift historical) → "
        "Pre-processing (clean, English-filter, deduplicate) → "
        "Trust-Score Module (rule-based heuristics on metadata + retail-insider terminology) → "
        "Model Analysis — sentiment with fine-tuned ModernBERT, aspects with zero-shot "
        "DeBERTa-v3, vision captioning with Gemma 3 4B (multi-pass prompt pipeline) when an "
        "image is present → "
        "Storage (SQLite analyses table) → "
        "Aggregation Layer (per-aspect / per-window summaries) → "
        "FastAPI REST + WebSocket → "
        "React Dashboard (Brand Health, Priority Negatives P1/P2, Review Queue, Pipeline panel). "
        "Reply drafts for P1/P2 posts are produced by Mistral 7B Instruct via Ollama and shown "
        "to reviewers; replies are never auto-posted.",
        italic=True)

    add_heading(doc, "Figure 2: End-to-End Pipeline Flow", level=3, bookmark="fig_2")
    add_image(doc, "fig2_pipeline_flow.png",
              "Figure 2: Single pipeline tick — 11 stages from subreddit discovery to dashboard "
              "surfacing, colour-coded by layer.", width_in=5.4)

    add_para(doc,
        "Design choices made during 11–30 May 2026:", bold=True, space_after=4)
    add_bullet(doc, "Data source: public Reddit only; no internal Walmart data is used.")
    add_bullet(doc, "Model stack — every stage uses a different model, chosen for the strengths "
                    "of that stage rather than one generic LLM:")
    add_bullet(doc, "   • Sentiment: a fine-tuned ModernBERT-base classifier (3-stage curriculum: "
                    "TweetEval → GoEmotions → Walmart-200 5-fold CV). Beats the cardiffnlp "
                    "twitter-roberta-base-sentiment-latest baseline by +0.137 macro-F1 overall "
                    "and +0.722 macro-F1 on long posts (≥512 tokens).")
    add_bullet(doc, "   • Aspect tagging: MoritzLaurer/deberta-v3-base-zeroshot-v2.0 used as a "
                    "zero-shot NLI classifier over a fixed 8-aspect label set (pricing, product "
                    "quality, customer service, store experience, online/app, delivery/pickup, "
                    "returns, app_website). No fine-tuning needed; multi-label, min_score=0.30.")
    add_bullet(doc, "   • Vision: Google Gemma 3 4B served locally via Ollama. Driven by a "
                    "multi-pass prompt pipeline (image-structure classification → per-tile OCR "
                    "→ merge to 2–4 sentences) designed from a 5-paper literature review on "
                    "fine-grained text recognition in vision-language models — UReader, "
                    "TextMonkey, DocOwl 1.5, InternVL2, Qwen2.5-VL.")
    add_bullet(doc, "   • Reply drafter (human-in-the-loop): Mistral 7B Instruct via Ollama; "
                    "upgrade path to llama3.1:8b is one config line.")
    add_bullet(doc, "   • Credibility / trust: rule-based heuristics in src/trust/heuristics.py "
                    "— promo-phrase regex, retail-insider terminology boost, and account "
                    "metadata weights. No LLM call is made for trust in the default local "
                    "pipeline; the cloud-LLM credibility path exists but is off by default to "
                    "keep cost and data-egress at zero.")
    add_bullet(doc, "Trust-score formula (Figure 3): "
                    "trust_score = 0.4·metadata_score + 0.3·dedup_originality "
                    "+ 0.3·heuristic_credibility, all components clipped to [0, 1]. The "
                    "metadata score itself is a weighted sum of account age, karma, post "
                    "length, engagement, and a base floor so brand-new short posts are not "
                    "unfairly penalised.")
    add_bullet(doc, "Storage: SQLite (data/local.db) with one row per post and one row per "
                    "analysis, indexed by post_id, subreddit, and created_ts for fast windowed "
                    "queries.")
    add_bullet(doc, "Scheduler: APScheduler running every 60 minutes (configurable), with "
                    "on-demand 'Run Now' and 'Backfill N days' controls exposed in the dashboard.")
    add_bullet(doc, "Dashboard: React + Vite + Tailwind, served at :3001, talking to FastAPI on "
                    ":8001. Walmart brand palette is used for visual identification only.")

    add_heading(doc, "Figure 3: Trust-Score Composition", level=3, bookmark="fig_3")
    add_image(doc, "fig3_trust_composition.png",
              "Figure 3: Trust-score inputs feed a rule-based heuristic scorer; the resulting "
              "trust_score is multiplied by sentiment_confidence to produce priority_score, "
              "which gates P1 and P2 tiers.")
    add_para(doc,
        "Components (default weights from config/pipeline_config.yaml): metadata_score (0.40) "
        "is a weighted sum of account age, karma, post length, engagement, and a base floor; "
        "dedup_originality (0.30) penalises near-duplicate text against the recent corpus using "
        "MinHash; heuristic_credibility (0.30) is a rule-based scorer that rewards "
        "retail-insider terminology (e.g. OGP, ASM, TLE, Spark, store #####) and penalises "
        "promotional / bot phrases via regex. The cloud-LLM credibility path exists in "
        "src/analysis/llm_client.py but is off in the default local pipeline.", italic=True)

    add_page_break(doc)

    # ===== 4. IMPLEMENTATION PROGRESS =====
    add_heading(doc, "4. Implementation Progress (25 April 2026 – 21 June 2026)", level=1, bookmark="sec_4")
    add_para(doc,
        "The four phases scheduled for the first half of the semester — literature review and "
        "outline, system design, data ingestion and pre-processing, and LLM analysis with "
        "trust scoring — have been completed. The implementation is committed to a private "
        "Walmart GitHub Enterprise repository "
        "(gecgithub01.walmart.com/v0s01jh/Retail_Sentiment_Intelligence) and runs end-to-end on "
        "a developer workstation.")

    # 4.1
    add_heading(doc, "4.1 Data Ingestion and Pre-processing", level=2, bookmark="sec_4_1")
    add_para(doc,
        "A dual-source ingestion module was implemented. PRAW (Python Reddit API Wrapper) is "
        "used for live polling of the curated subreddit list, and the public Arctic-Shift archive "
        "is used for historical backfill so that the prototype can be demonstrated against a "
        "non-empty corpus regardless of fresh activity. Per-subreddit cursors are persisted in "
        "SQLite so that incremental polls do not re-process posts.")
    add_para(doc,
        "Pre-processing includes language detection (English-only), removal of bot / "
        "auto-moderator content, URL stripping, MinHash-based near-duplicate detection, and "
        "PII redaction (emails, phone numbers, usernames in the body).", space_after=4)
    add_caption(doc, "Table 1: Curated Subreddit Coverage (snapshot 05 May 2026)", bookmark="tbl_1")
    add_table(doc,
        headers=["Segment", "Representative subreddits", "Approx. subscribers"],
        rows=[
            ["Walmart core", "r/walmart, r/WalmartEmployees, r/samsclub, r/OGPBackroom",
             "554,700"],
            ["Spark / last-mile", "r/Sparkdriver, r/walmartogp, r/WalmartSparkDrivers",
             "126,657"],
            ["Walmart pharmacy", "r/walmart_RX", "9,511"],
            ["Walmart international", "r/Flipkart, r/WalmartCanada", "18,877"],
            ["Retail competitors (benchmark)",
             "r/Costco, r/Target, r/AmazonPrime", "2,002,752"],
            ["Last-mile competitors (benchmark)",
             "r/doordash_drivers, r/AmazonFlexDrivers, r/instacart", "667,872"],
            ["Retail / customer voice (benchmark)",
             "r/MaliciousCompliance, r/TalesFromRetail, r/CustomerService", "6,198,120"],
        ],
        col_widths=[1.5, 3.2, 1.4])

    # 4.2
    add_heading(doc, "4.2 Model Analysis Module (Sentiment, Aspects, Vision)", level=2, bookmark="sec_4_2")
    add_para(doc,
        "The analysis layer deliberately uses three different models, one per task, instead of "
        "a single generic LLM. This was the most consequential design decision of the "
        "mid-semester window: it lowered cost to zero, removed external data egress, and "
        "produced measurably better results on long Reddit posts.")
    add_para(doc,
        "Sentiment classification — fine-tuned ModernBERT-base.", bold=True, space_after=2)
    add_para(doc,
        "A three-stage curriculum was used: (i) supervised pre-finetune on TweetEval "
        "sentiment, (ii) intermediate task on GoEmotions reduced to a three-class label set, "
        "(iii) final fine-tune on a hand-labelled Walmart-200 corpus using 5-fold "
        "cross-validation. The resulting checkpoint (models/modernbert_walmart/final, loaded "
        "via HuggingFace transformers, max_length=1024) beats the cardiffnlp "
        "twitter-roberta-base-sentiment-latest baseline by +0.137 macro-F1 overall and "
        "+0.722 macro-F1 on long posts (≥512 tokens) where RoBERTa is forced to truncate — "
        "see docs/MODEL_COMPARISON.md. If the local checkpoint is missing, the loader falls "
        "back to the RoBERTa baseline so the pipeline still runs on a fresh clone.")
    add_para(doc,
        "Aspect tagging — zero-shot DeBERTa-v3.", bold=True, space_after=2)
    add_para(doc,
        "Aspects are predicted by MoritzLaurer/deberta-v3-base-zeroshot-v2.0 used as an NLI "
        "classifier over an eight-label set (pricing, product quality, customer service, "
        "store experience, online/app, delivery/pickup, returns, app_website). Multi-label "
        "prediction with min_score=0.30 and at most three aspects per post keeps the output "
        "sparse. Zero-shot avoids the fine-tuning step entirely while remaining swappable for "
        "a supervised tagger later. Fallback: facebook/bart-large-mnli.")
    add_para(doc,
        "Vision — Gemma 3 4B (Ollama) with a multi-pass prompt pipeline.",
        bold=True, space_after=2)
    add_para(doc,
        "Approximately 12% of negative posts in our corpus include an image (receipts, "
        "shelf photos, app screenshots). A single 'describe this image' prompt against "
        "Gemma 3 4B was insufficient for screenshots that contain UI text. After a "
        "five-paper literature review on fine-grained text recognition in vision-language "
        "models — UReader, TextMonkey, DocOwl 1.5, InternVL2, and Qwen2.5-VL — a three-pass "
        "pipeline was implemented in src/analysis/vision.py: (i) STRUCTURE_PROMPT "
        "classifies the image type, (ii) the image is tiled and TILE_TEXT_PROMPT extracts "
        "verbatim text from each tile, (iii) MERGE_PROMPT fuses the structure verdict and "
        "tile texts into 2–4 sentences. Gemma 3 4B was chosen because it is natively "
        "multimodal in Ollama, scores 83 on DocVQA versus LLaVA-1.5's 28, fits in 8 GB RAM, "
        "and is permissively licensed (Google). LLaVA-7B is configured as fallback only. "
        "Qwen2.5-VL was excluded by enterprise policy on China-origin model providers.")
    add_para(doc,
        "Reply drafting (human-in-the-loop) — Mistral 7B Instruct.",
        bold=True, space_after=2)
    add_para(doc,
        "For P1/P2 negative posts, a Mistral 7B Instruct model served via Ollama drafts an "
        "empathetic reply (temperature=0.55, max_tokens=220). A deterministic "
        "'smart-composer' produces a second draft so reviewers can compare and choose. "
        "Replies are never auto-posted: the Reddit OAuth surface is gated by a dry_run flag.")
    add_para(doc,
        "Output contract.", bold=True, space_after=2)
    add_para(doc,
        "Each pipeline tick emits a JSON record per post containing overall_sentiment in "
        "{positive, negative, neutral}, sentiment_confidence in [0, 1], up to three aspects "
        "with per-aspect score, optional vision_caption when an image was present, and the "
        "trust score and flags. The output is validated, repaired on parse failure, and "
        "persisted to the analyses table. Cost is tracked per record in data/llm_costs.jsonl "
        "(zero by default because every step runs locally).")

    add_caption(doc, "Table 2: Eight-Aspect Retail Taxonomy", bookmark="tbl_2")
    add_table(doc,
        headers=["Aspect (zero-shot label)", "Examples of posts that match"],
        rows=[
            ["pricing", "Price mismatch, promo not honoured, perceived price gouging."],
            ["product quality", "Defective items, expired groceries, private-brand quality."],
            ["customer service", "Store staff behaviour, escalations, chat / call experience."],
            ["store experience", "In-store hygiene, shelf availability, register queues."],
            ["online/app", "Walmart.com or Walmart app crash, checkout errors, search bugs."],
            ["delivery/pickup", "Late or wrong delivery, Spark driver behaviour, OGP pickup delays."],
            ["returns", "Refund delays, refund rejected, return-process complaints."],
            ["app_website", "Account / login / sign-up problems, password reset failures."],
        ],
        col_widths=[2.0, 4.0])

    # 4.3
    add_heading(doc, "4.3 Trust-Score Module", level=2, bookmark="sec_4_3")
    add_para(doc,
        "The trust-score module computes the composition described in Figure 3. Account-age, "
        "karma, post length, and engagement are read from the public Reddit user/post objects; "
        "the metadata sub-score is a weighted sum of those plus a base floor so brand-new short "
        "posts are not unfairly penalised. The dedup-originality sub-score uses MinHash "
        "similarity against the recent corpus to penalise copy-paste / templated spam. The "
        "credibility sub-score is a rule-based scorer in src/trust/heuristics.py that boosts "
        "posts containing Walmart-insider terminology (OGP, Spark, ASM, TLE, CSM, GM, Cap 2, "
        "OnePOS, GTA, store/aisle/department/register #####) and penalises promotional or bot "
        "phrases via regex. No LLM call is made on the default local path; a cloud-LLM "
        "credibility option exists but is disabled in pipeline_config.yaml. The final "
        "trust_score is persisted alongside each analysis so downstream queries can apply the "
        "threshold without recomputation.")

    # 4.4
    add_heading(doc, "4.4 Storage Layer", level=2, bookmark="sec_4_4")
    add_para(doc,
        "Posts and analyses are stored in a local SQLite database (data/local.db). The schema "
        "separates raw_posts (immutable ingest payload) from analyses (the LLM output plus "
        "trust score and validation flags). Indexes on (subreddit, created_ts) and "
        "(sentiment, trust_score) keep the dashboard windowed queries below ~50 ms on the "
        "current corpus.")

    # 4.5
    add_heading(doc, "4.5 Dashboard and Human-in-the-Loop Review", level=2, bookmark="sec_4_5")
    add_para(doc,
        "A React + Vite + Tailwind dashboard (5 pages) is served on :3001 and talks to a "
        "FastAPI backend on :8001. Figure 4 shows how the pages, global filters, storage, "
        "and reviewer feedback loop are wired together. Pages implemented during 01–21 June 2026:")
    add_bullet(doc, "Brand Health — KPI tiles (Total, Trusted, Positive, Negative, P1, P2), "
                    "sentiment pie chart, segment distribution, top issues, and a new "
                    "Priority Negative Posts panel that lists the top-N (10/15/20/30/50/100) "
                    "negative posts ranked by trust × confidence, tiered as P1 "
                    "(trust ≥ 0.70 AND confidence ≥ 0.80) and P2 "
                    "(trust ≥ 0.50 AND confidence ≥ 0.60).")
    add_bullet(doc, "Post Explorer — filterable post list by subreddit, sentiment, aspect, "
                    "time range, and minimum trust score.")
    add_bullet(doc, "Aspect Drilldown — per-aspect trends and representative example posts.")
    add_bullet(doc, "Review Queue — human-in-the-loop interface for low-confidence posts; "
                    "corrections are saved back to the analyses table and feed all dashboards.")
    add_bullet(doc, "Pipeline — live pipeline status with per-subreddit ingest progress, "
                    "media-capture coverage, manual 'Run Now' and 'Backfill' controls, and a "
                    "danger-zone reset.")

    add_heading(doc, "Figure 4: Dashboard Information Architecture", level=3, bookmark="fig_4")
    add_image(doc, "fig4_dashboard_map.png",
              "Figure 4: Dashboard information architecture — six pages backed by a single "
              "FastAPI surface, sharing global filters, and feeding reviewer corrections back "
              "into the storage layer.")

    add_page_break(doc)

    # ===== 5. TOOLS / TECHNOLOGIES =====
    add_heading(doc, "5. Tools, Technologies, and Configuration", level=1, bookmark="sec_5")
    add_para(doc,
        "Table 3 lists the tools and technologies used in the prototype. All components run "
        "locally on a developer workstation; no Walmart-internal infrastructure is used.")
    add_caption(doc, "Table 3: Tools and Technologies Used", bookmark="tbl_3")
    add_table(doc,
        headers=["Layer", "Tool / Library", "Purpose"],
        rows=[
            ["Language", "Python 3.13", "Backend, ingestion, analysis, scheduler."],
            ["Reddit ingestion", "PRAW", "Live polling of curated subreddits."],
            ["Historical backfill", "Arctic-Shift API",
             "Public Reddit archive for non-empty demo corpora."],
            ["Scheduler", "APScheduler (60-min interval)",
             "Background polling, with on-demand 'Run Now' and 'Backfill N days' triggers."],
            ["Sentiment model", "ModernBERT-base, fine-tuned (HuggingFace transformers)",
             "3-stage curriculum (TweetEval → GoEmotions → Walmart-200 5-fold CV); +0.137 macro-F1 over RoBERTa baseline."],
            ["Sentiment fallback", "cardiffnlp/twitter-roberta-base-sentiment-latest",
             "Used automatically when the fine-tuned checkpoint is absent."],
            ["Aspect tagging", "MoritzLaurer/deberta-v3-base-zeroshot-v2.0",
             "Zero-shot NLI over an 8-label retail taxonomy; multi-label, min_score=0.30."],
            ["Vision", "Gemma 3 4B via Ollama (multi-pass prompt pipeline)",
             "Native-multimodal captioning + tile-OCR; design informed by UReader / TextMonkey / DocOwl 1.5 / InternVL2 / Qwen2.5-VL."],
            ["Reply drafter", "Mistral 7B Instruct via Ollama",
             "Human-in-the-loop reply drafts for P1 / P2 posts; never auto-posted."],
            ["Trust / credibility", "Rule-based heuristics (regex + retail-insider terms)",
             "src/trust/heuristics.py; no LLM call on default local path."],
            ["Storage", "SQLite (data/local.db)",
             "raw_posts and analyses tables, plus cursors and reviewer corrections."],
            ["API", "FastAPI + Uvicorn", "REST endpoints and WebSocket for live status."],
            ["Frontend", "React + Vite + TypeScript + Tailwind",
             "Six-page dashboard (Brand Health, Aspect, Post Explorer, Alert Feed, Review, Pipeline)."],
            ["Charts", "Recharts", "Sentiment pie, time-series trends, aspect breakdowns."],
            ["Privacy", "Custom regex redactor",
             "Strips emails, phones, and usernames from stored bodies."],
        ],
        col_widths=[1.3, 1.9, 2.8])

    # ===== 6. PROBLEMS & MITIGATIONS =====
    add_heading(doc, "6. Problems Encountered and Mitigations", level=1, bookmark="sec_6")
    add_para(doc,
        "Several non-trivial blockers were hit while building the prototype. The four largest "
        "(data access, model access, vision quality, and trust formula) reshaped key design "
        "decisions; the remainder were operational issues fixed in passing. Each row in "
        "Table 4 captures both the original plan and the working alternative that replaced it.")
    add_caption(doc, "Table 4: Problems Encountered and Mitigations", bookmark="tbl_4")
    add_table(doc,
        headers=["#", "Problem", "Mitigation"],
        rows=[
            ["1",
             "Reddit data access — the original plan was to ingest live posts via PRAW with "
             "developer-API credentials. The Reddit developer-API application was submitted but "
             "approval did not arrive within the mid-semester window, so PRAW could not be "
             "authenticated against a real client_id / client_secret. Without credentials, PRAW "
             "falls back to anonymous read which is heavily rate-limited (60 req/min) and "
             "returns a very thin recent stream — the demo dashboard would be empty for hours.",
             "Switched the default ingestion path to the public Arctic-Shift archive "
             "(arctic-shift.photon-reddit.com), which exposes Reddit's historical corpus over "
             "a free, credential-less HTTPS API. The fetcher_provider toggle in "
             "config/pipeline_config.yaml lets us flip back to PRAW the moment the dev-API key "
             "arrives. A 'Backfill N days' control in the dashboard now lets the demo run "
             "deterministically against any historical window."],
            ["2",
             "Foundation-model access — the abstract called for GPT-4o (Azure OpenAI) for "
             "sentiment, aspects, and credibility. The Azure OpenAI subscription / API-key "
             "provisioning ticket was raised but the key was not issued in time, and Walmart "
             "data-egress rules also discourage routing public-Reddit text through external "
             "hosted endpoints without a documented approval.",
             "Re-architected the analysis layer around in-house, locally served open-weight "
             "models. Three different models, one per task, were selected for the strengths of "
             "each stage: a fine-tuned ModernBERT-base for sentiment (HuggingFace transformers), "
             "DeBERTa-v3 zero-shot NLI for aspects, and Gemma 3 4B served via Ollama for "
             "vision. Mistral 7B Instruct (Ollama) drafts reviewer replies. Cost dropped to "
             "zero, no data leaves the workstation, and the LLMClient interface keeps the "
             "Azure OpenAI path available behind a one-line provider switch in "
             "config/models.yaml when the key eventually arrives."],
            ["3",
             "Vision / image-post processing — a single-shot 'describe this image' prompt "
             "against Gemma 3 4B was insufficient for the dominant image type in our corpus "
             "(app screenshots, receipts, shelf-tag photos containing UI text). The model "
             "described elements in isolation ('a blue button', 'Learn More text') without "
             "connecting them into 'the customer's checkout failed at confirmation', and it "
             "missed verbatim price / error-code text on receipts.",
             "Reviewed five recent papers on fine-grained text recognition in vision-language "
             "models — UReader, TextMonkey, DocOwl 1.5, InternVL2, Qwen2.5-VL — and built a "
             "three-pass prompt pipeline in src/analysis/vision.py instead of changing the "
             "model: (i) STRUCTURE_PROMPT classifies image type, (ii) the image is tiled and "
             "TILE_TEXT_PROMPT extracts verbatim text from each tile, (iii) MERGE_PROMPT fuses "
             "structure + tile texts into a 2–4 sentence caption. This recovered the missing "
             "UI text without exceeding the 8 GB RAM budget and without switching to "
             "Qwen2.5-VL (excluded by enterprise policy on China-origin providers)."],
            ["4",
             "Working trust-and-credibility formula — the first attempt routed every post "
             "through an LLM credibility check, which doubled cost (later: zero, but added "
             "latency on local models) and was overkill: an account with 1,200 days of age, "
             "5,600 karma, and Walmart-insider terminology in the body should not need a model "
             "call to be trusted, while a 1-day-old account posting promo URLs should not need "
             "one to be flagged. A flat weighted sum of all heuristic signals also produced too "
             "many borderline cases for reviewers, and dropping low-trust posts (the initial "
             "design) destroyed downstream recall.",
             "Iterated to a three-component formula "
             "(trust_score = 0.4·metadata + 0.3·dedup + 0.3·heuristic_credibility) in "
             "src/trust/scorer.py. The metadata sub-score is itself a weighted sum of account "
             "age, karma, post length, engagement, and a base floor so brand-new short posts "
             "are not unfairly penalised. The credibility sub-score is rule-based: regex over "
             "promotional / bot phrases, plus a positive boost for retail-insider terms "
             "(OGP, ASM, TLE, CSM, Spark, OnePOS, store/aisle/department #####). Low-trust "
             "posts are now flagged, not dropped (per R5 in the requirements). The optional "
             "LLM credibility path is retained and only invoked when 0.3 < metadata_score < 0.8 "
             "— i.e. only the genuinely ambiguous band."],
            ["5",
             "Sentiment-model selection and training — the initial supervised baseline used a "
             "smaller pre-trained encoder (the cardiffnlp twitter-roberta-base-sentiment-latest "
             "checkpoint, ~125M params, 512-token cap). It performed reasonably on tweet-length "
             "text but degraded sharply on long Reddit posts (≥512 tokens) where it had to "
             "truncate the middle of the complaint; it also under-weighted retail jargon it had "
             "never seen during pre-training.",
             "Migrated to ModernBERT-base (newer encoder, 8K-token context, better long-range "
             "attention) and trained a 3-stage curriculum: (i) supervised pre-finetune on "
             "TweetEval sentiment to anchor the label space, (ii) intermediate task on "
             "GoEmotions reduced to a three-class mapping for richer affect coverage, "
             "(iii) final fine-tune on a hand-labelled Walmart-200 corpus using 5-fold "
             "cross-validation. The resulting checkpoint (models/modernbert_walmart/final, "
             "max_length=1024) beats the RoBERTa baseline by +0.137 macro-F1 overall and "
             "+0.722 macro-F1 on long posts — see docs/MODEL_COMPARISON.md. RoBERTa is kept as "
             "an automatic fallback if the checkpoint is absent on a fresh clone."],
            ["6",
             "Raw model output was occasionally non-JSON — trailing prose, half-closed "
             "objects, or extra commentary — which broke the persistence step and lost the "
             "whole record.",
             "Added a JSON-repair layer: extract the first balanced JSON object, validate "
             "against the analysis schema, and on failure re-prompt with a tighter system "
             "message; the post is skipped (with a logged reason) only after two retries so a "
             "single bad response cannot poison the batch."],
            ["7",
             "Triage volume — once sentiment + trust were live, reviewers asked for a clear "
             "'what should I look at first' view; a flat sorted feed of negative posts was too "
             "long to act on.",
             "Added a Priority Negatives panel on Brand Health that ranks negative posts by "
             "priority_score = trust_score × sentiment_confidence and tiers them into P1 "
             "(trust ≥ 0.70 AND confidence ≥ 0.80) and P2 (trust ≥ 0.50 AND confidence ≥ "
             "0.60). The thresholds are shown inline on the tiles and on the section header so "
             "the rule is not hidden in code."],
            ["8",
             "Pipeline opacity — long backfills (90 days × 32 subreddits) ran for many "
             "minutes and operators could not tell whether the pipeline was progressing or "
             "stuck.",
             "Added subreddit_fetch_start / progress / complete events from the ingestor and "
             "surfaced a live collapsible per-subreddit progress panel on the Pipeline page, "
             "plus a status pill in the header that mirrors the FastAPI /api/pipeline/status "
             "snapshot."],
            ["9",
             "Media-capture coverage briefly read >100% because the denominator counted only "
             "posts whose source was a single image, while the numerator also captioned "
             "videos, GIFs (animated reddit_video / preview/.gif assets that the vision step "
             "successfully caption-keyframes), and link-card previews.",
             "Re-defined images_total as max(image_only + text_plus_image + video, captioned) "
             "so the percentage is clamped to ≤100% by construction. The KPI tile now also "
             "shows the raw 'captioned / images_total' fraction in a tooltip so the operator "
             "can see what was counted."],
            ["10",
             "Banned / restricted / empty subreddits in the curated list (r/PhonePe, r/Equate, "
             "r/walmartspark, r/walmartmexico) caused the ingestor to raise 403 / 404 mid-run "
             "and abort the whole batch.",
             "Annotated the subreddit registry with a status field "
             "(public / restricted / banned / empty); the ingestor now skips non-public "
             "entries with a logged warning and continues with the remaining list."],
        ],
        col_widths=[0.4, 2.8, 2.8])

    add_page_break(doc)

    # ===== 7. PRELIMINARY OBSERVATIONS =====
    add_heading(doc, "7. Preliminary Observations", level=1, bookmark="sec_7")
    add_para(doc,
        "Formal evaluation is scheduled for 01–10 July 2026 against a manually labelled sample. "
        "Preliminary qualitative observations from spot-checking the running pipeline as of "
        "21 June 2026 are summarised below; these are not evaluation results.")
    add_bullet(doc, "The fine-tuned ModernBERT classifier handles negation, sarcasm, and "
                    "retail-specific slang (e.g., 'OGP pickup', 'ETL', 'TL', 'Spark') noticeably "
                    "better than the cardiffnlp RoBERTa baseline, particularly on long posts "
                    "where the older model is forced to truncate.")
    add_bullet(doc, "Aspect tagging is most reliable for delivery, returns, and pricing; "
                    "customer support and app/website experience occasionally overlap and will "
                    "need a clarifying prompt example.")
    add_bullet(doc, "Trust filtering visibly removes a small but non-trivial set of promotional / "
                    "low-effort posts that would otherwise inflate the negative bucket.")
    add_bullet(doc, "The Priority Negatives panel surfaces credible, actionable complaints "
                    "(e.g., r/Target store-staff escalations) that would have been buried in a "
                    "time-sorted feed.")
    add_bullet(doc, "End-to-end ingest → analyse → display latency for a single post is in the "
                    "low seconds on the local Ollama setup, dominated by LLM inference rather "
                    "than ingestion or storage.")

    # ===== 8. FUTURE PLAN =====
    add_heading(doc, "8. Future Plan (Remaining Phases)", level=1, bookmark="sec_8")
    add_para(doc,
        "The phases below mirror the plan submitted with the abstract. Status as of "
        "21 June 2026 is shown in the rightmost column.")
    add_caption(doc, "Table 5: Future Plan with Status", bookmark="tbl_5")
    add_table(doc,
        headers=["Phase", "Start Date – End Date", "Work to be done", "Status"],
        rows=[
            ["Outline & Literature Review",
             "25 Apr 2026 – 10 May 2026",
             "Literature review on LLM-based sentiment / aspect analysis and fake-review / "
             "credibility detection. Submit dissertation outline.",
             "COMPLETED"],
            ["Design",
             "11 May 2026 – 30 May 2026",
             "Finalize data sources, per-task model choices (ModernBERT for sentiment, "
             "DeBERTa-v3 zero-shot for aspects, Gemma 3 4B for vision, Mistral 7B for reply "
             "drafts), aspect taxonomy, trust-score design, evaluation plan, and overall "
             "system design.",
             "COMPLETED"],
            ["Data Ingestion & Pre-processing",
             "11 May 2026 – 30 May 2026",
             "Implement Reddit ingestion (PRAW + Arctic-Shift), cleaning, English filtering, "
             "deduplication, and storage. Collect a baseline working corpus.",
             "COMPLETED"],
            ["Model Development & Trust Scoring",
             "01 Jun 2026 – 21 Jun 2026",
             "Fine-tune ModernBERT on Walmart-200 (3-stage curriculum), wire DeBERTa-v3 "
             "zero-shot for aspects, design the multi-pass Gemma 3 4B vision pipeline from a "
             "5-paper literature review, and implement the rule-based trust scorer. Integrate "
             "with the dashboard.",
             "COMPLETED"],
            ["Aggregation & Reporting",
             "22 Jun 2026 – 30 Jun 2026",
             "Aggregated summaries per aspect and time window. Notebook-based reporting and "
             "Brand Health dashboard refinements (Priority Negatives P1/P2, drill-downs).",
             "IN PROGRESS"],
            ["Evaluation",
             "01 Jul 2026 – 10 Jul 2026",
             "Manually label ~250–300 posts. Measure sentiment accuracy / macro-F1, aspect "
             "macro-F1, and ROC-AUC for the trust filter on a small balanced credibility sample.",
             "PENDING"],
            ["Dissertation Review",
             "11 Jul 2026 – 20 Jul 2026",
             "Submit draft to Supervisor and Additional Examiner; incorporate feedback.",
             "PENDING"],
            ["Final Submission",
             "21 Jul 2026 – 02 Aug 2026",
             "Final review, formatting per WILP guidelines, and submission.",
             "PENDING"],
        ],
        col_widths=[1.6, 1.5, 2.6, 1.0])

    add_page_break(doc)

    # ===== 9. ABBREVIATIONS =====
    add_heading(doc, "9. Abbreviations", level=1, bookmark="sec_9")
    add_table(doc,
        headers=["Abbreviation", "Expansion"],
        rows=[
            ["API", "Application Programming Interface"],
            ["APScheduler", "Advanced Python Scheduler"],
            ["AUC", "Area Under the (ROC) Curve"],
            ["BITS", "Birla Institute of Technology and Science"],
            ["BLIP", "Bootstrapping Language-Image Pre-training"],
            ["F1", "Harmonic mean of precision and recall"],
            ["GIF", "Global Integrated Fulfillment"],
            ["HITL", "Human-in-the-Loop"],
            ["JSON", "JavaScript Object Notation"],
            ["LLaVA", "Large Language-and-Vision Assistant"],
            ["LLM", "Large Language Model"],
            ["LVDS", "(not used in this report)"],
            ["MinHash", "Probabilistic near-duplicate similarity hash"],
            ["MTTR", "Mean Time To Resolve"],
            ["NLP", "Natural Language Processing"],
            ["OGP", "Online Grocery Pickup"],
            ["OAuth", "Open Authorization"],
            ["OCR", "Optical Character Recognition"],
            ["PRAW", "Python Reddit API Wrapper"],
            ["P1 / P2", "Priority Tier 1 / Tier 2 (priority_score-ranked negatives)"],
            ["ROC", "Receiver Operating Characteristic"],
            ["REST", "Representational State Transfer"],
            ["SLA", "Service Level Agreement"],
            ["SQL", "Structured Query Language"],
            ["SQLite", "Embedded SQL database engine"],
            ["WILP", "Work-Integrated Learning Programmes (BITS Pilani)"],
        ],
        col_widths=[1.6, 4.4])

    # ===== 10. REFERENCES =====
    add_heading(doc, "10. References", level=1, bookmark="sec_10")
    refs = [
        "[1] DataReportal (We Are Social & Meltwater). \"Digital 2024: Global Overview Report,\" January 2024.",
        "[2] Reddit, Inc. Form S-1 Registration Statement filed with the U.S. SEC, February 2024.",
        "[3] Sprout Social. \"The Sprout Social Index 2024,\" 2024.",
        "[4] World Economic Forum. \"Fake online reviews are a $152 billion problem. Here's how to silence them,\" 2024.",
        "[5] T. Brown, B. Mann, N. Ryder, et al. \"Language Models are Few-Shot Learners.\" NeurIPS, vol. 33, pp. 1877–1901, 2020.",
        "[6] OpenAI. \"GPT-4 Technical Report.\" arXiv:2303.08774, 2023.",
        "[7] D. Zhang, S. Li, H. Zhu, et al. \"Sentiment Analysis of Social Media Using Large Language Models: A Comprehensive Survey.\" arXiv:2311.10066, 2023.",
        "[8] K. Shu, A. Sliva, S. Wang, J. Tang, and H. Liu. \"Fake News Detection on Social Media: A Data Mining Perspective.\" ACM SIGKDD Explorations Newsletter, 19(1):22–36, 2017.",
        "[9] A. Mukherjee, B. Liu, and N. Glance. \"Spotting Fake Reviewer Groups in Consumer Reviews.\" WWW, pp. 191–200, 2012.",
        "[10] Walmart Inc. \"Corporate Fact Sheet\" (FY2024).",
        "[11] B. Liu. Sentiment Analysis and Opinion Mining. Morgan & Claypool, 2012.",
        "[12] H. Liu, I. Chatterjee, M. Zhou, X. S. Lu, and A. Srivastava. \"Aspect-Based Sentiment Analysis: A Survey of Deep Learning Methods.\" IEEE TCSS, 10(1):211–232, 2023.",
        "[13] J. Wei, X. Wang, D. Schuurmans, et al. \"Chain-of-Thought Prompting Elicits Reasoning in Large Language Models.\" NeurIPS, vol. 35, 2022.",
        "[14] J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova. \"BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding.\" NAACL-HLT, pp. 4171–4186, 2019.",
        "[15] H. Touvron, L. Martin, K. Stone, et al. \"Llama 2: Open Foundation and Fine-Tuned Chat Models.\" arXiv:2307.09288, 2023.",
    ]
    for r in refs:
        add_para(doc, r, size=10, space_after=2)

    # save
    doc.save(OUT_PATH)
    print(f"Wrote: {OUT_PATH}")


if __name__ == "__main__":
    build()
