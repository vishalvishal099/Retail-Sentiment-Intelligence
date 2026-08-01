"""
Build the BITS WILP FINAL Dissertation Report (.docx) for:
    Vishal Singh, 2020AA05641
    Real-Time Social Media Mining and Trust-Aware Sentiment Analysis
    Using Large Language Models for Retail Product Feedback Optimization

Structure follows the BITS "Guidelines for preparation of WILP Project Report"
(06.04.2022) required elements: Cover, Title Page, Acknowledgements, Abstract
Sheet, Table of Contents, Introduction, Main Text (design, implementation,
evaluation, results), Conclusions & Recommendations, Appendices, References,
Glossary. Formatting helpers mirror the mid-sem builder
(docs/Sem_4/build_mid_sem_report.py).

Run with the conda python that has python-docx:
    /opt/miniconda3/bin/python docs/Sem_4/final/build_final_report.py

Output:
    docs/Sem_4/final/FINAL_REPORT_VishalSingh_2020AA05641.docx
"""

from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT_PATH = Path(__file__).parent / "FINAL_REPORT_VishalSingh_2020AA05641.docx"
FIGURES_DIR = Path(__file__).parent.parent / "figures"

# ── Typography (thesis-style, LaTeX book-class-like) ──
BODY_FONT = "Times New Roman"
HEAD_FONT = "Times New Roman"
BODY_SIZE = 12         # main text
BODY_LEADING = 1.5     # 1.5 line spacing
H1_SIZE = 17           # section headings inside chapters (e.g. 1.1)
H2_SIZE = 14           # subsection (e.g. 5.1.1)
H3_SIZE = 12           # sub-subsection
CHAP_LABEL_SIZE = 24   # 'Chapter N' label
CHAP_TITLE_SIZE = 24   # chapter title, same as label
CAPTION_SIZE = 11
TABLE_BODY_SIZE = 11
TABLE_HEAD_SIZE = 11
NAVY_RGB = RGBColor(0x04, 0x1E, 0x42)

# LaTeX book class inhibits first-line indent immediately after a heading, then
# indents every subsequent paragraph.  We emulate that behaviour with a flag.
_JUST_HEADED = [False]


# ───────────────────────── helpers ─────────────────────────

def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=110, right=110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def _set_outline_level(paragraph, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


_BOOKMARK_ID = [2000]


def _add_bookmark(paragraph, name: str) -> None:
    bm_id = str(_BOOKMARK_ID[0])
    _BOOKMARK_ID[0] += 1
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bm_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bm_id)
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is not None:
        p_pr.addnext(start)
    else:
        paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_toc_link(doc, text: str, bookmark: str, indent: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.25
    if indent:
        pf.left_indent = Inches(0.35)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rFont = OxmlElement("w:rFonts")
    rFont.set(qn("w:ascii"), BODY_FONT)
    rFont.set(qn("w:hAnsi"), BODY_FONT)
    rpr.append(rFont)
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    rpr.append(color)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(size * 2))
    rpr.append(sz)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single")
    rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t"); t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    hyperlink.append(r)
    p._p.append(hyperlink)


def add_caption(doc, text: str, bookmark: str | None = None) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.2
    pf.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = HEAD_FONT
    run.font.size = Pt(CAPTION_SIZE)
    run.font.color.rgb = RGBColor(0x04, 0x1E, 0x42)
    if bookmark:
        _add_bookmark(p, bookmark)


def add_heading(doc, text: str, level: int = 1, bookmark: str | None = None) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(20 if level == 1 else 14 if level == 2 else 10)
    pf.space_after = Pt(10 if level == 1 else 6)
    pf.keep_with_next = True
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = True
    run.font.name = HEAD_FONT
    if level == 1:
        run.font.size = Pt(H1_SIZE)
    elif level == 2:
        run.font.size = Pt(H2_SIZE)
    else:
        run.font.size = Pt(H3_SIZE)
    run.font.color.rgb = NAVY_RGB
    _set_outline_level(p, max(0, level - 1))
    if bookmark:
        _add_bookmark(p, bookmark)
    _JUST_HEADED[0] = True


def add_para(doc, text: str, bold: bool = False, italic: bool = False,
             size: int | float | None = None, align=None, space_after: int = 0,
             justify: bool = True, first_line_indent: bool = True) -> None:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = BODY_LEADING
    apply_indent = first_line_indent and align is None and justify and not _JUST_HEADED[0]
    if apply_indent:
        pf.first_line_indent = Inches(0.28)
    if align is not None:
        p.alignment = align
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE if size is None else size)
    run.bold = bold
    run.italic = italic
    if text.strip():
        _JUST_HEADED[0] = False


def add_rich_para(doc, segments: list[tuple], size: int | float | None = None,
                  align=None, space_after: int = 0, justify: bool = True,
                  first_line_indent: bool = True) -> None:
    """Add a paragraph containing multiple runs.
    Each segment is (text, bold, italic) or (text, bold).  Empty segments are skipped."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = BODY_LEADING
    apply_indent = first_line_indent and align is None and justify and not _JUST_HEADED[0]
    if apply_indent:
        pf.first_line_indent = Inches(0.28)
    if align is not None:
        p.alignment = align
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fsize = Pt(BODY_SIZE if size is None else size)
    for seg in segments:
        if not seg or not seg[0]:
            continue
        text = seg[0]
        bold = seg[1] if len(seg) > 1 else False
        italic = seg[2] if len(seg) > 2 else False
        run = p.add_run(text)
        run.font.name = BODY_FONT
        run.font.size = fsize
        run.bold = bold
        run.italic = italic
    _JUST_HEADED[0] = False


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.4
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)


def add_numbered(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    pf = p.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.4
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)


def add_table(doc, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None, header_bg: str = "041E42",
              zebra: str = "F2F4F7") -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(h)
        run.bold = True
        run.font.name = HEAD_FONT
        run.font.size = Pt(TABLE_HEAD_SIZE)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr_cells[i], header_bg)
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=110, right=110)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        zebra_on = (r_idx % 2 == 0)
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.18
            run = p.add_run(str(val))
            run.font.name = BODY_FONT
            run.font.size = Pt(TABLE_BODY_SIZE)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cells[c_idx], top=70, bottom=70, left=110, right=110)
            if zebra_on:
                set_cell_bg(cells[c_idx], zebra)
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(w)


def add_image(doc, filename: str, caption: str, width_in: float = 6.3) -> None:
    path = FIGURES_DIR / filename
    if not path.exists():
        add_para(doc, f"[Missing figure: {filename} — run docs/Sem_4/generate_diagrams.py]",
                 italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(2)
    pf.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_in))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(14)
    cap.paragraph_format.line_spacing = 1.15
    crun = cap.add_run(caption)
    crun.italic = True
    crun.font.name = HEAD_FONT
    crun.font.size = Pt(CAPTION_SIZE)
    crun.font.color.rgb = RGBColor(0x04, 0x1E, 0x42)


def add_page_break(doc) -> None:
    doc.add_page_break()


def set_page_numbering(section, fmt: str = "decimal", start: int | None = None) -> None:
    """Control page-number format per section: 'lowerRoman' (i, ii) or 'decimal' (1, 2)."""
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def setup_page_chrome(doc) -> None:
    for section in doc.sections:
        # A4 (like Gobind's LaTeX report)
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)
        hp = section.header.paragraphs[0]
        hp.text = ""
        fp = section.footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        inner = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), BODY_FONT)
        rFonts.set(qn("w:hAnsi"), BODY_FONT)
        rpr.append(rFonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22")
        rpr.append(sz)
        inner.append(rpr)
        txt = OxmlElement("w:t"); txt.text = "1"
        inner.append(txt)
        fld.append(inner)
        fp._p.append(fld)


# ───────────────────────── chapter helper ─────────────────────────

_CHAP_NO = [0]


def add_chapter(doc, title: str, bookmark: str | None = None) -> None:
    """Start a new page and render a LaTeX book-class-style chapter header:
       'Chapter N' and the chapter title on two separate lines at the same 24pt bold."""
    _CHAP_NO[0] += 1
    doc.add_page_break()
    # push chapter title ~1.4 inches down the page like LaTeX book class
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for _ in range(6):
        spacer_run = spacer.add_run("\n")
        spacer_run.font.size = Pt(12)
    lab = doc.add_paragraph()
    lab.paragraph_format.space_before = Pt(0)
    lab.paragraph_format.space_after = Pt(12)
    lab.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    lab.paragraph_format.first_line_indent = Inches(0)
    r = lab.add_run(f"Chapter {_CHAP_NO[0]}")
    r.bold = True
    r.font.name = HEAD_FONT
    r.font.size = Pt(CHAP_LABEL_SIZE)
    r.font.color.rgb = NAVY_RGB
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(30)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run(title)
    run.bold = True
    run.font.name = HEAD_FONT
    run.font.size = Pt(CHAP_TITLE_SIZE)
    run.font.color.rgb = NAVY_RGB
    _set_outline_level(p, 0)
    if bookmark:
        _add_bookmark(p, bookmark)


# ───────────────────────── build ─────────────────────────

def build() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_SIZE)
    pf = style.paragraph_format
    pf.line_spacing = BODY_LEADING
    pf.space_after = Pt(6)
    # Ensure East-Asian font slot also uses Times New Roman so Word applies it
    from docx.oxml.ns import qn as _qn
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(_qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(_qn("w:ascii"), BODY_FONT)
    rFonts.set(_qn("w:hAnsi"), BODY_FONT)
    rFonts.set(_qn("w:cs"), BODY_FONT)

    for section in doc.sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.0)
    setup_page_chrome(doc)
    # Front matter uses lower-case roman page numbers (i, ii, iii, ...)
    set_page_numbering(doc.sections[0], fmt="lowerRoman", start=1)

    C = WD_ALIGN_PARAGRAPH.CENTER
    R = WD_ALIGN_PARAGRAPH.RIGHT
    TITLE = ("REAL-TIME SOCIAL MEDIA MINING AND TRUST-AWARE SENTIMENT ANALYSIS "
             "USING LARGE LANGUAGE MODELS FOR RETAIL PRODUCT FEEDBACK OPTIMIZATION")

    # ══════════════ (i) COVER — Appendix A (minimal outer cover) ══════════════
    add_para(doc, "", space_after=72)
    add_para(doc, "A REPORT", bold=True, size=16, align=C, space_after=12)
    add_para(doc, "ON", bold=True, size=13, align=C, space_after=12)
    add_para(doc, TITLE, bold=True, size=16, align=C, space_after=72)
    add_para(doc, "BY", bold=True, size=13, align=C, space_after=12)
    add_para(doc, "VISHAL SINGH", bold=True, size=15, align=C, space_after=2)
    add_para(doc, "ID No. 2020AA05641", size=12, align=C, space_after=72)
    add_para(doc, "Walmart Global Tech, Bengaluru, India", bold=True, size=13, align=C, space_after=48)
    add_para(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI", bold=True, size=13, align=C, space_after=2)
    add_para(doc, "November 2026", size=12, align=C, space_after=2)
    add_page_break(doc)

    # ══════════════ (ii) TITLE PAGE — Appendix B (detailed inner title page) ══════════════
    add_para(doc, "", space_after=30)
    add_para(doc, "A REPORT", bold=True, size=15, align=C, space_after=10)
    add_para(doc, "ON", bold=True, size=13, align=C, space_after=10)
    add_para(doc, TITLE, bold=True, size=15, align=C, space_after=30)
    add_para(doc, "BY", bold=True, size=13, align=C, space_after=10)
    add_para(doc, "Vishal Singh", bold=True, size=14, align=C, space_after=2)
    add_para(doc, "ID No. 2020AA05641", size=12, align=C, space_after=2)
    add_para(doc, "Discipline: M.Tech (Artificial Intelligence & Machine Learning)", size=12, align=C, space_after=24)
    add_para(doc, "Prepared in partial fulfilment of the", size=11, align=C, space_after=2)
    add_para(doc, "WILP Dissertation Course  (BITS ZG628T)", bold=True, size=12, align=C, space_after=24)
    add_para(doc, "Under the Supervision of", size=11, align=C, space_after=2)
    add_para(doc, "Mr. Varunendra Pratap Singh", bold=True, size=12, align=C, space_after=2)
    add_para(doc, "Principal Software Engineer, Walmart Global Tech, Bengaluru", size=11, align=C, space_after=24)
    add_para(doc, "AT", bold=True, size=13, align=C, space_after=10)
    add_para(doc, "Walmart Global Tech, Bengaluru, Karnataka, India", bold=True, size=12, align=C, space_after=24)
    add_para(doc, "BIRLA INSTITUTE OF TECHNOLOGY & SCIENCE, PILANI (RAJASTHAN)", bold=True, size=13, align=C, space_after=2)
    add_para(doc, "Work Integrated Learning Programmes Division", size=11, align=C, space_after=2)
    add_para(doc, "November 2026", size=12, align=C, space_after=2)
    add_page_break(doc)

    # ══════════════ CERTIFICATE (from Supervisor) ══════════════
    add_heading(doc, "Certificate", level=1, bookmark="sec_cert")
    add_para(doc, "", space_after=6)
    add_para(doc,
        "This is to certify that the Dissertation report entitled \u201c" + TITLE + "\u201d "
        "submitted by Mr. Vishal Singh (ID No. 2020AA05641) in partial fulfilment of the "
        "requirements of the course BITS ZG628T Dissertation, embodies the work done by him "
        "under my supervision and guidance.", space_after=18)
    add_para(doc,
        "The work is original and, to the best of my knowledge, has not been submitted "
        "elsewhere for the award of any other degree or diploma. During the period of the "
        "dissertation the candidate has shown sound technical competence, initiative, and a "
        "commitment to an honest, defensible treatment of results.", space_after=36)
    cert = doc.add_table(rows=2, cols=2)
    cert.autofit = True
    cert.cell(0, 0).text = ""
    cert.cell(0, 1).text = "Signature of the Supervisor"
    cert.cell(1, 0).text = "Place: Bengaluru\nDate:"
    cert.cell(1, 1).text = ("Name: Mr. Varunendra Pratap Singh\n"
                            "Designation: Principal Software Engineer\n"
                            "Walmart Global Tech, Bengaluru")
    for row in cert.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    add_page_break(doc)

    # ══════════════ (iii) ACKNOWLEDGEMENTS ══════════════
    add_heading(doc, "Acknowledgements", level=1, bookmark="sec_ack")
    add_para(doc,
        "I take this opportunity to thank everyone who supported me during the course of this "
        "dissertation.")
    add_rich_para(doc, [
        ("I am grateful to the leadership at ", False),
        ("Walmart Global Tech India", True),
        (" for providing an environment that encouraged learning and experimentation, and "
         "for allowing me to pursue this industry-linked academic project alongside my regular "
         "engineering responsibilities.", False),
    ])
    add_rich_para(doc, [
        ("I sincerely thank my Supervisor, ", False),
        ("Mr. Varunendra Pratap Singh", True),
        (" (Principal Software Engineer, Walmart Global Tech), for his day-to-day technical "
         "guidance, patient reviews of my design and evaluation choices, and for continually "
         "pushing me to raise the standard of the work. I also thank the Additional Examiner, "
         "assigned by the organisation, for the rigorous mid-project scrutiny that "
         "substantially improved the evaluation methodology and the trust-score design.", False),
    ])
    add_para(doc,
        "I thank my colleagues and professional experts on the retail domain at Walmart Global "
        "Tech for helping me translate free-form Reddit posts into an operational, "
        "eight-aspect retail taxonomy that drives the aggregation and alerting stages of this "
        "system.")
    add_rich_para(doc, [
        ("I am deeply grateful to my Faculty Mentor, ", False),
        ("Ms. Pradnya Kashikar", True),
        (" (BITS Pilani, WILP Division), for her guidance, timely feedback, and academic "
         "oversight from the abstract-outline stage through mid-semester review to final "
         "submission, and for ensuring the work conformed to the academic rigour expected of "
         "an M.Tech dissertation.", False),
    ])
    add_para(doc,
        "Finally, I thank my family and friends for their patience and encouragement over the "
        "course of this programme, and I dedicate this effort to them.")
    add_para(doc,
        "All experiments reported in this dissertation use only publicly available data; no "
        "proprietary Walmart data or internal systems were used at any stage.")
    add_para(doc, "", space_after=18)
    add_para(doc, "Vishal Singh", bold=True, size=12, justify=False, space_after=0)
    add_para(doc, "Bengaluru, November 2026", size=12, justify=False, space_after=18)
    sig = doc.add_table(rows=2, cols=1)
    sig.autofit = True
    sig.cell(0, 0).text = "Signature of the Student"
    sig.cell(1, 0).text = "Name: Vishal Singh    ID No.: 2020AA05641\nDate:    Place: Bengaluru"
    for row in sig.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    add_page_break(doc)

    # ══════════════ (iv) ABSTRACT SHEET — Appendix C ══════════════
    add_para(doc, "BIRLA INSTITUTE OF TECHNOLOGY AND SCIENCE, PILANI (RAJASTHAN)",
             bold=True, size=12, align=C, space_after=2)
    add_para(doc, "Work Integrated Learning Programmes Division", size=11, align=C, space_after=2)
    add_para(doc, "Dissertation Abstract Sheet", bold=True, size=12, align=C, space_after=10)
    absheet = [
        ["Organization", "Walmart Global Tech"],
        ["Location", "Bengaluru, Karnataka, India"],
        ["Duration", "Two semesters (approximately 8 months)"],
        ["Date of Start", "25 April 2026"],
        ["Date of Submission", "November 2026"],
        ["Title of the Project", TITLE],
        ["ID No. / Name of the Student", "2020AA05641 / Vishal Singh"],
        ["Name & Designation of Supervisor and Additional Examiner",
         "Mr. Varunendra Pratap Singh, Principal Software Engineer, Walmart Global Tech "
         "(Supervisor); Additional Examiner as assigned by the organisation"],
        ["Name of the Faculty Mentor", "Ms. Pradnya Kashikar, BITS Pilani WILP Division"],
        ["Key Words",
         "Sentiment Analysis; Aspect-Based Opinion Mining; Large Language Models; ModernBERT; "
         "Zero-shot NLI; Trust / Credibility Filtering; Multimodal Vision; Reddit; Retail Analytics"],
        ["Project Areas",
         "Applied Natural Language Processing; Machine Learning; Multimodal AI; Data Engineering"],
    ]
    add_table(doc, ["Field", "Detail"], absheet, col_widths=[2.1, 4.2])
    add_para(doc, "Abstract", bold=True, size=11, space_after=4)
    add_para(doc,
        "Public social media is now a primary real-time channel for retail customer voice, yet "
        "fake and low-credibility posts contaminate any aggregate consumed without filtering. "
        "This dissertation designs, builds, and evaluates Retail Sentiment Intelligence (RSI), a "
        "complete offline-first prototype that converts a noisy public Reddit stream into a "
        "structured, aspect-tagged, trust-weighted brand-health feed. Posts from 25 curated "
        "retail communities are ingested, de-duplicated, and assigned a 0\u20131 trust score "
        "combining metadata heuristics, near-duplicate detection, and credibility checks. Trusted "
        "posts are classified for sentiment by a domain fine-tuned ModernBERT encoder, tagged "
        "with aspects by a zero-shot NLI classifier over an eight-category taxonomy, and \u2014 "
        "when an image is attached \u2014 captioned by a multi-pass, anti-hallucination Gemma 3 4B "
        "pipeline. Sentiment macro-F1 improved from 0.6272 to 0.7642 (and 0.28 to 1.00 on long "
        "posts); vision hallucination fell from 50% to 0%. A React/FastAPI dashboard surfaces "
        "results across eight surfaces including a live alert feed, a human-in-the-loop review "
        "queue, a post-lifecycle board, and competitor insights. The report documents the full "
        "lifecycle: literature survey, design, implementation, honest cross-validated evaluation, "
        "limitations, and recommendations.", space_after=18)
    asig = doc.add_table(rows=2, cols=2)
    asig.autofit = True
    asig.cell(0, 0).text = "Signature of the Student"
    asig.cell(0, 1).text = "Signature of the Supervisor"
    asig.cell(1, 0).text = "Name: Vishal Singh\nDate:"
    asig.cell(1, 1).text = "Name: Varunendra Pratap Singh\nDate:"
    for row in asig.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    add_page_break(doc)

    # ══════════════ (v) TABLE OF CONTENTS ══════════════
    add_heading(doc, "Table of Contents", level=1, bookmark="sec_toc")
    add_para(doc,
        "Front matter is numbered in lower-case Roman numerals (i, ii, iii, …). The main text is "
        "numbered in Arabic numerals starting at page 1 from the Introduction. Section numbers "
        "below are clickable and jump to the corresponding heading.",
        italic=True, size=9.5, space_after=8)
    toc = [
        ("Certificate", "sec_cert", False),
        ("Acknowledgements", "sec_ack", False),
        ("Abstract", "sec_abstract", False),
        ("Chapter 1  Introduction", "sec_1", False),
        ("1.1  Broad Area of Work", "sec_1_1", True),
        ("1.2  Motivation", "sec_1_2", True),
        ("1.3  Contributions of this Dissertation", "sec_1_3", True),
        ("1.4  Organisation of the Report", "sec_1_4", True),
        ("Chapter 2  Problem Statement and Objectives", "sec_2", False),
        ("Chapter 3  Literature Survey", "sec_3", False),
        ("Chapter 4  System Design and Architecture", "sec_4", False),
        ("4.1  Design Principles", "sec_4_1", True),
        ("4.2  Layered Architecture", "sec_4_2", True),
        ("4.3  End-to-End Pipeline Flow", "sec_4_3", True),
        ("Chapter 5  Implementation", "sec_5", False),
        ("5.1  Data Ingestion and Pre-processing", "sec_5_1", True),
        ("5.2  Trust-Score Module", "sec_5_2", True),
        ("5.3  Sentiment Classification (ModernBERT)", "sec_5_3", True),
        ("5.4  Aspect-Based Opinion Mining", "sec_5_4", True),
        ("5.5  Vision / Image Processing", "sec_5_5", True),
        ("5.6  Aggregation, Alerts, and Notifications", "sec_5_6", True),
        ("5.7  Dashboard Overview", "sec_5_7", True),
        ("5.8  Human-in-the-Loop Review & Validate", "sec_5_8", True),
        ("5.9  Smart Reply Composer — Dual-Draft (Rule + LLM)", "sec_5_9", True),
        ("5.10  Learning Loop — Feedback for Future Retraining", "sec_5_10", True),
        ("5.11  Post Explorer and Multi-Facet Search", "sec_5_11", True),
        ("5.12  Post Lifecycle (Kanban Workflow)", "sec_5_12", True),
        ("5.13  Insights and Competitor Analysis", "sec_5_13", True),
        ("5.14  Storage Layer", "sec_5_14", True),
        ("Chapter 6  Evaluation and Results", "sec_6", False),
        ("6.1  Sentiment Model Evaluation", "sec_6_1", True),
        ("6.2  Vision Module Evaluation", "sec_6_2", True),
        ("6.3  Trust-Score Behaviour", "sec_6_3", True),
        ("Chapter 7  Tools, Technologies, and Configuration", "sec_7", False),
        ("Chapter 8  Problems Encountered and Mitigations", "sec_8", False),
        ("Chapter 9  Conclusions and Recommendations", "sec_9", False),
        ("Chapter 10  Future Work", "sec_10", False),
        ("Glossary and Abbreviations", "sec_11", False),
        ("References", "sec_12", False),
        ("Appendix A — Curated Subreddit Coverage", "app_a", False),
        ("Appendix B — Reproduction Commands", "app_b", False),
        ("Checklist of Items for the Final Report", "sec_checklist", False),
    ]
    for text, bm, indent in toc:
        add_toc_link(doc, text, bm, indent=indent)
    add_para(doc, "", space_after=8)
    add_para(doc, "List of Figures", bold=True, size=11, space_after=4)
    figs = [
        ("Figure 1: Layered System Architecture", "fig_1"),
        ("Figure 2: End-to-End Pipeline Flow", "fig_2"),
        ("Figure 3: Trust-Score Composition and P1/P2 Tier Rules", "fig_3"),
        ("Figure 4: Dashboard Information Architecture and Feedback Loop", "fig_4"),
        ("Figure 5: Human-in-the-Loop Review & Validate Flow", "fig_5"),
        ("Figure 6: Smart Reply Composer — Dual-Draft with Few-Shot", "fig_6"),
        ("Figure 7: Feedback Learning Loop and Path to Auto-Reply", "fig_7"),
        ("Figure 8: Post Lifecycle Kanban Workflow", "fig_8"),
        ("Figure 9: Data Ingestion and Pre-processing Flow", "fig_9"),
        ("Figure 10: ModernBERT Three-Stage Fine-Tuning Curriculum", "fig_10"),
        ("Figure 11: Zero-Shot NLI Aspect Tagging", "fig_11"),
        ("Figure 12: Vision Multi-Pass Anti-Hallucination Pipeline", "fig_12"),
        ("Figure 13: Alert Engine and Notification Routing", "fig_13"),
        ("Figure 14: Sentiment Macro-F1 — RoBERTa vs ModernBERT", "fig_14"),
    ]
    for text, bm in figs:
        add_toc_link(doc, text, bm)
    add_para(doc, "", space_after=8)
    add_para(doc, "List of Tables", bold=True, size=11, space_after=4)
    tbls = [
        ("Table 1: Eight-Aspect Retail Taxonomy", "tbl_1"),
        ("Table 2: Sentiment Model Comparison (out-of-fold CV)", "tbl_2"),
        ("Table 3: Per-Length-Bucket Sentiment F1", "tbl_3"),
        ("Table 4: Vision Module — Before vs After Multi-Pass", "tbl_4"),
        ("Table 5: Tools and Technologies", "tbl_5"),
        ("Table 6: Problems Encountered and Mitigations", "tbl_6"),
        ("Table 7: Curated Subreddit Coverage", "tbl_7"),
    ]
    for text, bm in tbls:
        add_toc_link(doc, text, bm)

    # ── Section break: main matter restarts at Arabic page 1 ──
    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    main_section.page_height = Inches(11.69)
    main_section.page_width = Inches(8.27)
    main_section.top_margin = Inches(1.0)
    main_section.bottom_margin = Inches(1.0)
    main_section.left_margin = Inches(1.25)
    main_section.right_margin = Inches(1.0)
    set_page_numbering(main_section, fmt="decimal", start=1)

    # One-page standalone Abstract (Introduction begins on page 1 after this)
    add_heading(doc, "Abstract", level=1, bookmark="sec_abstract")
    add_para(doc,
        "Public social media is now a primary real-time channel for retail customer voice, yet "
        "fake and low-credibility posts contaminate any aggregate consumed without filtering. "
        "This dissertation designs, builds, and evaluates Retail Sentiment Intelligence (RSI), a "
        "complete offline-first prototype that converts a noisy public Reddit stream into a "
        "structured, aspect-tagged, trust-weighted brand-health feed. Posts from 25 curated "
        "retail communities are ingested, de-duplicated, and assigned a 0\u20131 trust score "
        "combining metadata heuristics, near-duplicate detection, and credibility checks. Trusted "
        "posts are classified for sentiment by a domain fine-tuned ModernBERT encoder, tagged "
        "with aspects by a zero-shot NLI classifier over an eight-category taxonomy, and \u2014 "
        "when an image is attached \u2014 captioned by a multi-pass, anti-hallucination Gemma 3 4B "
        "pipeline. Sentiment macro-F1 improved from 0.6272 to 0.7642 (and 0.28 to 1.00 on long "
        "posts); vision hallucination fell from 50% to 0%. A React/FastAPI dashboard surfaces "
        "results across eight surfaces including a live alert feed, a human-in-the-loop review "
        "queue, a post-lifecycle board, and competitor insights. The report documents the full "
        "lifecycle: literature survey, design, implementation, honest cross-validated evaluation, "
        "limitations, and recommendations.")
    add_para(doc,
        "Key Words: Sentiment Analysis; Aspect-Based Opinion Mining; Large Language Models; "
        "ModernBERT; Zero-shot NLI; Trust / Credibility Filtering; Multimodal Vision; Reddit; "
        "Retail Analytics.", bold=True, size=10, space_after=4)
    add_para(doc,
        "Project Areas: Applied Natural Language Processing; Machine Learning; Multimodal AI; "
        "Data Engineering.", bold=True, size=10)
    add_page_break(doc)


    # ══════════════ 1. INTRODUCTION ══════════════
    add_chapter(doc, "Introduction", bookmark="sec_1")

    add_heading(doc, "1.1  Broad Area of Work", level=2, bookmark="sec_1_1")
    add_para(doc,
        "The broad area of this dissertation is Applied Natural Language Processing (NLP) and "
        "multimodal machine learning for the analysis of public retail customer feedback on "
        "social media. The work spans four technical sub-areas: (a) supervised text "
        "classification with transformer encoders for sentiment in short, informal English "
        "social-media text; (b) zero-shot natural-language-inference (NLI) models for "
        "aspect-based opinion mining over a fixed retail taxonomy; (c) multimodal "
        "vision-language models for captioning and OCR of screenshots and photographs attached "
        "to retail complaints; and (d) lightweight trust / credibility filtering using "
        "account-metadata heuristics, near-duplicate detection, and retail-insider terminology "
        "signals. Around these models sits a data-engineering layer for scheduled ingestion of "
        "public posts and structured storage suitable for aggregation and reporting, and a "
        "serving layer that surfaces the results in an interactive dashboard.")

    add_heading(doc, "1.2  Motivation", level=2, bookmark="sec_1_2")
    add_para(doc,
        "The motivation is volume and noise. Customers, associates, and delivery drivers post "
        "continuously about their Walmart, Sam's Club, Spark-driver, and OGP/pickup experiences "
        "on Reddit. In the candidate's role at Walmart Global Tech, this signal is presently "
        "consumed through ad-hoc browsing and keyword-filter dashboards, which suffer from poor "
        "coverage, high noise, and latency: a complaint that is trending today may not surface in "
        "a manual review for days. Conventional lexicon- or rule-based sentiment systems perform "
        "poorly on the sarcasm, slang, and long-form narrative style typical of Reddit retail "
        "posts, and none of them filter fake or low-credibility content before aggregation. This "
        "dissertation demonstrates that a stack of task-specific models — a fine-tuned encoder "
        "for sentiment, a zero-shot NLI model for aspects, and a multimodal vision model for "
        "images — combined with an explicit trust filter can convert a raw public stream into a "
        "structured, aspect-tagged, trust-weighted feed suitable for operational review.")

    add_heading(doc, "1.3  Contributions of this Dissertation", level=2, bookmark="sec_1_3")
    add_para(doc, "The principal contributions of the completed work are:", bold=True)
    add_numbered(doc, "A modular, offline-first, zero-API-cost end-to-end pipeline (Retail "
                      "Sentiment Intelligence) covering ingestion, trust scoring, sentiment, "
                      "aspects, vision, aggregation, alerting, and dashboarding.")
    add_numbered(doc, "A domain fine-tuned ModernBERT sentiment classifier trained through a "
                      "three-stage curriculum, raising macro-F1 from 0.6272 to 0.7642 overall and "
                      "recovering all long posts (≥ 512 tokens, n=7, all negative-class) that the "
                      "RoBERTa baseline mis-classifies (5/7 → 7/7 correct), evaluated with honest "
                      "5-fold out-of-fold cross-validation.")
    add_numbered(doc, "A multi-pass vision-captioning technique that adapts ideas from recent "
                      "(policy-blocked) vision-language papers onto a compliant Gemma 3 4B model, "
                      "reducing hallucination from 50% to 0% and lifting text extraction from "
                      "25% to 75%.")
    add_numbered(doc, "An interpretable trust score and a trust × confidence admission gate that "
                      "flags — rather than drops — low-credibility posts, with every constant "
                      "mapped to a stakeholder-arguable English rationale.")
    add_numbered(doc, "A React/FastAPI dashboard (Brand Health, Alert Feed, Post Explorer, "
                      "Review & Validate, Post Lifecycle, Insights & Competitor, Pipeline Control, "
                      "Notifications) with a human-in-the-loop Review & Validate workflow whose "
                      "corrections and posted replies feed few-shot reply generation and a "
                      "roadmap to automatic replies, closing the loop between AI output and "
                      "analyst action.")

    add_heading(doc, "1.4  Organisation of the Report", level=2, bookmark="sec_1_4")
    add_para(doc,
        "The remainder of this report is organised into nine further chapters. Chapter 2 states "
        "the problem precisely and sets out the objectives and scope of the dissertation. "
        "Chapter 3 surveys the relevant literature on transformer sentiment classification, "
        "aspect-based opinion mining, multimodal vision-language models, and social-media "
        "credibility. Chapter 4 presents the layered system design and the end-to-end pipeline "
        "flow. Chapter 5 describes the implementation of each module in detail — ingestion, "
        "trust scoring, sentiment, aspects, vision, aggregation and alerting, the dashboard, and "
        "the human-in-the-loop and learning-loop workflows. Chapter 6 reports the experimental "
        "evaluation and results for the sentiment, vision, and trust components. Chapter 7 "
        "documents the tools, technologies, and configuration used. Chapter 8 records the "
        "problems encountered during the work and the mitigations adopted. Chapter 9 draws "
        "conclusions and recommendations, and Chapter 10 outlines directions for future work. "
        "A glossary, the references, and supporting appendices close the report.")

    # ══════════════ 2. PROBLEM STATEMENT ══════════════
    add_chapter(doc, "Problem Statement and Objectives", bookmark="sec_2")
    add_para(doc,
        "Manual monitoring of retail-related Reddit communities cannot keep pace with the daily "
        "volume of posts, and conventional rule-based or lexicon-based sentiment systems perform "
        "poorly on short, sarcastic, slang-heavy retail text. Long-form complaints that carry the "
        "decisive detail (a receipt, a store timeline, a resolution) are silently truncated by "
        "512-token models. Furthermore, fake or low-credibility posts contaminate any aggregate "
        "signal if consumed without filtering. The problem is therefore to build a system that "
        "(i) collects retail social posts at scale, (ii) filters low-credibility content "
        "explicitly, (iii) classifies sentiment and aspects accurately on long-form domain text, "
        "(iv) handles image-only posts without fabricating detail, and (v) presents the result in "
        "a form an analyst can act on.")
    add_para(doc, "The objectives approved for the dissertation were:", bold=True)
    add_numbered(doc, "Survey recent literature on LLM-based sentiment and aspect-based opinion "
                      "mining, and on fake-review / bot-credibility detection.")
    add_numbered(doc, "Build a data-ingestion pipeline that periodically collects retail-related "
                      "posts from Reddit via the official Reddit API and a public historical archive.")
    add_numbered(doc, "Implement a model-based analysis module that assigns each post a sentiment "
                      "label and one or more aspects from a fixed retail taxonomy, plus a vision "
                      "caption when an image is attached.")
    add_numbered(doc, "Implement a trust score (0–1) per post combining account-metadata "
                      "heuristics, near-duplicate detection, and a rule-based credibility scorer "
                      "(with an optional cloud-LLM credibility path).")
    add_numbered(doc, "Produce structured aggregated outputs (top aspects, sentiment distribution "
                      "per aspect, representative examples, priority-negatives ranking) with "
                      "low-trust posts flagged for review.")
    add_numbered(doc, "Evaluate the prototype on a manually labelled sample of 200 posts: "
                      "sentiment macro-F1 (target ≥ 0.80 as a stretch gate), aspect coverage, and "
                      "trust-filter behaviour.")
    add_numbered(doc, "Document design, evaluation, limitations, and recommended next steps in the "
                      "final dissertation.")

    # ══════════════ 3. LITERATURE SURVEY ══════════════
    add_chapter(doc, "Literature Survey", bookmark="sec_3")
    add_para(doc,
        "The survey covered three strands. First, transformer encoders for sentiment: RoBERTa "
        "fine-tuned on TweetEval [2] remains a strong public baseline for short social text, but "
        "is capped at 512 tokens and trained on tweets rather than long-form Reddit narratives. "
        "ModernBERT [1] is a modern BERT-style encoder with an 8,192-token context and a "
        "web/code/long-document training corpus, making it a natural candidate for long retail "
        "complaints while retaining BERT-class inference throughput. GoEmotions [3] supplies a "
        "large corpus of Reddit comments used to adapt the model to the Reddit register.")
    add_para(doc,
        "Second, aspect-based opinion mining: rather than training a supervised ABSA model on a "
        "small labelled set, the project uses zero-shot NLI classification [4] with a DeBERTa-v3 "
        "[5] entailment model, which frames each candidate aspect as a hypothesis to be entailed "
        "by the post. This requires no training data, supports multi-aspect assignment, and lets "
        "the taxonomy be edited without retraining — a good fit for a domain whose aspect "
        "vocabulary (pricing, delivery/pickup, returns, app/website, etc.) is known a priori.")
    add_para(doc,
        "Third, vision-language models for document/screenshot understanding. Five recent papers "
        "— UReader [7], TextMonkey [8], DocOwl 1.5 [9], InternVL2 [10], and Qwen2.5-VL [11] "
        "(2023–2025) — converge on a single conclusion: dynamic / native-resolution processing "
        "with tile-based attention is the key to fine-grained text recognition, and fixed-size "
        "image resizing is the dominant failure mode. Crucially, every model that implements this "
        "natively is from a lab whose provider is blocked by enterprise vendor policy. The design "
        "response — detailed in Section 5.5 — was to re-implement those techniques as a calling "
        "strategy on a compliant Gemma 3 4B model [6], with de-duplication driven by MiniLM "
        "sentence embeddings [12].")
    add_para(doc,
        "Finally, on credibility: the fake-review literature and its economic impact [13] "
        "motivate a multi-signal filter (account age, karma, posting frequency, duplicate "
        "detection, promotional-language heuristics) rather than a single learned classifier, "
        "which would be hard to justify on only 200 labelled posts. The project therefore adopts "
        "an interpretable, weighted rule-based trust score. Market context for the problem is "
        "drawn from industry reports on social usage and expected response times [14][15][16].")

    # ══════════════ 4. SYSTEM DESIGN ══════════════
    add_chapter(doc, "System Design and Architecture", bookmark="sec_4")

    add_heading(doc, "4.1  Design Principles", level=2, bookmark="sec_4_1")
    add_para(doc, "Four principles governed every design decision:", bold=True)
    add_bullet(doc, "Local-first and free by default — every model runs offline (HuggingFace / "
                    "Ollama); cloud LLMs and paid APIs are gated behind explicit keys, and "
                    "per-call cost is tracked in a JSONL ledger.")
    add_bullet(doc, "Modularity — each pipeline stage is independently testable and swappable "
                    "(e.g., SQLite ↔ Cosmos DB, RoBERTa ↔ ModernBERT) via configuration.")
    add_bullet(doc, "Flag, don't drop — low-trust posts are surfaced for human review, never "
                    "silently discarded, preserving recall and auditability.")
    add_bullet(doc, "Defensibility — evaluation uses honest out-of-fold cross-validation, not "
                    "training-set fit; every trust constant maps to an arguable English sentence.")

    add_heading(doc, "4.2  Layered Architecture", level=2, bookmark="sec_4_2")
    add_para(doc,
        "The system is organised as six logical layers. Figure 1 shows the layers and the "
        "components inside each.")
    add_caption(doc, "Figure 1: Layered System Architecture", bookmark="fig_1")
    add_image(doc, "fig1_architecture.png",
              "Figure 1: Six-layer architecture — data sources, ingestion/processing pipeline, "
              "AI runtime, storage, serving (FastAPI + WebSocket), and clients (dashboard, "
              "email, Slack).")
    add_para(doc,
        "In words: Reddit sources (Arctic Shift historical API + optional PRAW live) feed an "
        "asyncio ingestion and processing pipeline that cleans, English-filters, and "
        "de-duplicates posts, then trust-scores them and runs the AI-analysis stage. The AI "
        "runtime hosts ModernBERT (sentiment), a zero-shot NLI classifier (aspects), Gemma 3 4B "
        "(vision, via Ollama), and FLAN-T5 / Mistral (reply drafting). Results persist to six "
        "tables (SQLite in development, Azure Cosmos DB in production). A FastAPI serving layer "
        "exposes a REST API and a WebSocket alert channel to a React dashboard, and dispatches "
        "email/Slack notifications for priority posts.",
        italic=True)

    add_heading(doc, "4.3  End-to-End Pipeline Flow", level=2, bookmark="sec_4_3")
    add_para(doc,
        "A single scheduler tick (default cadence: every 6 hours, plus manual 'Run Now') passes "
        "each new post through six stages: Ingest → Pre-process → Trust-Score → Analyze "
        "(sentiment + aspects + optional vision) → Aggregate → Alert. Figure 2 traces this flow.")
    add_caption(doc, "Figure 2: End-to-End Pipeline Flow", bookmark="fig_2")
    add_image(doc, "fig2_pipeline_flow.png",
              "Figure 2: Single pipeline tick — six stages from cursor-based ingestion of 25 "
              "subreddits to dashboard surfacing and notification dispatch.", width_in=5.4)
    add_para(doc,
        "Total latency is roughly 3–5 minutes for a full sweep of 25 subreddits. Ingestion is "
        "incremental and cursor-based, so only posts created since the last tick are fetched, "
        "except for an initial back-fill of up to 90 days per community on first run.")

    # ══════════════ 5. IMPLEMENTATION ══════════════
    add_chapter(doc, "Implementation", bookmark="sec_5")

    add_heading(doc, "5.1  Data Ingestion and Pre-processing", level=2, bookmark="sec_5_1")
    add_para(doc,
        "Posts are ingested from 25 curated communities spanning Walmart core, Spark drivers, "
        "Sam's Club, pharmacy, international, OGP/pickup, and retail competitors (full list in "
        "Appendix A). The primary source is the free Arctic Shift historical API with cursor-based "
        "pagination; PRAW (the official Reddit API) is an optional live source. Pre-processing "
        "strips HTML/Markdown, applies a langdetect English filter, and de-duplicates using "
        "MiniLM-L6-v2 sentence embeddings with a cosine-similarity threshold of 0.92. Reddit "
        "usernames are SHA-hashed before storage, and a one-year retention default applies, "
        "satisfying the privacy requirements.")
    add_caption(doc, "Figure 9: Data Ingestion and Pre-processing Flow", bookmark="fig_9")
    add_image(doc, "fig9_ingestion_flow.png",
              "Figure 9: The ingestion path — 25 subreddits via Arctic Shift / PRAW, cleaning, "
              "English filtering, MiniLM de-duplication, username hashing, into the raw_posts "
              "table.", width_in=6.2)

    add_heading(doc, "5.2  Trust-Score Module", level=2, bookmark="sec_5_2")
    add_para(doc,
        "Every post receives a trust score in [0, 1] as a weighted combination of three "
        "orthogonal signals:")
    add_para(doc,
        "trust_score = 0.4 · meta + 0.3 · dedup + 0.3 · llm   (clipped to [0, 1])",
        bold=True, align=C)
    add_para(doc,
        "The metadata sub-score (weight 0.4) rewards account age, karma, content length, and "
        "engagement, each capped at a saturation point (1 year, 5,000 karma, 200 characters, "
        "score 20) above a 0.15 floor. The dedup sub-score (0.3) is 1.0 unless a near-duplicate "
        "exists, in which case it drops to 0.5. The LLM/rule credibility sub-score (0.3) is only "
        "invoked in the ambiguous zone (0.3 < meta < 0.8) for cost control, and checks for "
        "promotional language, URL stuffing, ALL-CAPS, karma/age mismatch, and — as positive "
        "signals — retail-insider terminology (OGP, ASM, CAP2, Spark) and long-form organic text.")
    add_para(doc,
        "A post is admitted to influence the dashboard only when trust_score × sentiment_"
        "confidence ≥ 0.30. Multiplication (not addition) is deliberate: if either factor is near "
        "zero the gate closes, so a confident label cannot rescue an untrustworthy post, nor can "
        "high trust rescue a genuinely ambiguous sentiment call. Figure 3 shows the composition "
        "and the P1/P2 priority tiers.")
    add_caption(doc, "Figure 3: Trust-Score Composition and P1/P2 Tier Rules", bookmark="fig_3")
    add_image(doc, "fig3_trust_composition.png",
              "Figure 3: Trust-score composition (0.4 meta + 0.3 dedup + 0.3 llm), the "
              "trust × confidence admission gate at 0.30, and the P1/P2 priority-negative tiers.")
    add_para(doc,
        "Priority tiers for negative posts are: P1 when trust ≥ 0.70 and confidence ≥ 0.80 "
        "(high-signal, immediate action); P2 when trust ≥ 0.50 and confidence ≥ 0.60 "
        "(review-worthy, lower urgency). These tiers drive both the dashboard priority-negatives "
        "panel and the group-based notification routing.")

    add_heading(doc, "5.3  Sentiment Classification (ModernBERT)", level=2, bookmark="sec_5_3")
    add_para(doc,
        "Sentiment is a 3-class problem (positive / negative / neutral). The production model is "
        "a fine-tuned answerdotai/ModernBERT-base, chosen over the RoBERTa baseline for its "
        "8,192-token context (16× RoBERTa's 512) and long-document training corpus. Fine-tuning "
        "used a three-stage curriculum: Stage 1 on TweetEval-sentiment (45k tweets) for generic "
        "grounding; Stage 2 on GoEmotions-3class (54k Reddit comments) for Reddit register; and "
        "Stage 3 on the 200 hand-labelled Walmart posts with stratified 5-fold cross-validation, "
        "class weighting (neg=0.52, neu=1.03, pos=8.33), minority oversampling, and a "
        "1,024-token max length. Reported numbers are out-of-fold — every sample is predicted by "
        "a fold model that never trained on it. A separate final model, trained on all 200 posts "
        "with the same recipe, is the artifact the production pipeline loads; it is never used to "
        "compute reported metrics.")
    add_caption(doc, "Figure 10: ModernBERT Three-Stage Fine-Tuning Curriculum", bookmark="fig_10")
    add_image(doc, "fig10_modernbert_curriculum.png",
              "Figure 10: The three-stage curriculum — TweetEval grounding → GoEmotions Reddit "
              "register → Walmart-200 domain specialisation → deployed final model, with per-stage "
              "out-of-fold macro-F1.", width_in=6.2)

    add_heading(doc, "5.4  Aspect-Based Opinion Mining", level=2, bookmark="sec_5_4")
    add_para(doc,
        "Aspects are assigned by zero-shot NLI classification over the eight-category taxonomy in "
        "Table 1. Each candidate aspect is framed as a hypothesis; the post entails zero or more "
        "of them (multi-aspect is allowed). This requires no training data and lets the taxonomy "
        "evolve without retraining. Customer and employee aspects are kept in separate "
        "sub-taxonomies so that, for example, an associate's pay-and-benefits complaint is not "
        "mislabelled as a customer pricing issue.")
    add_caption(doc, "Figure 11: Zero-Shot NLI Aspect Tagging", bookmark="fig_11")
    add_image(doc, "fig11_aspect_nli.png",
              "Figure 11: Each candidate aspect is framed as a hypothesis and tested for "
              "entailment by DeBERTa-v3; aspects scoring above threshold are assigned, allowing "
              "multi-aspect labels with no training data.", width_in=6.2)
    add_caption(doc, "Table 1: Eight-Aspect Retail Taxonomy", bookmark="tbl_1")
    add_table(doc,
        ["#", "Aspect", "Scope / Examples"],
        [
            ["1", "pricing", "Prices, promotions, price-match, perceived value"],
            ["2", "product_quality", "Damaged/defective goods, freshness, wrong item"],
            ["3", "customer_support", "Store or phone service, associate behaviour, resolution"],
            ["4", "store_experience", "In-store layout, cleanliness, checkout, stock-outs"],
            ["5", "online_app", "Website / app usability, bugs, login, account issues"],
            ["6", "delivery_pickup", "OGP, curbside, Spark delivery, substitutions, ETA"],
            ["7", "returns", "Return / refund policy and process friction"],
            ["8", "workforce_hr", "(Employee) pay, benefits, management, workload, safety"],
        ],
        col_widths=[0.5, 1.8, 4.0])

    add_heading(doc, "5.5  Vision / Image Processing", level=2, bookmark="sec_5_5")
    add_para(doc,
        "About 3.9% of Reddit posts have empty bodies where the complaint lives entirely inside "
        "an attached image (an app error, a damaged product, a receipt). A text-only pipeline "
        "misses these. The chosen vision model is Gemma 3 4B via Ollama — the best DocVQA score "
        "(83) under 4 GB, Google-maintained (policy-compliant), and reusing existing Ollama "
        "infrastructure. Single-pass captioning failed on 75% of an initial 8-image test and "
        "hallucinated on 50% (fabricating fake receipts and prices), which would corrupt the "
        "dashboard.")
    add_para(doc,
        "The mitigation is a four-pass pipeline that borrows techniques from the surveyed "
        "(policy-blocked) papers: Pass 1 (structure) asks what type the image is; Pass 2 (tiling) "
        "splits it into 2–4 crops for 2–4× effective resolution; Pass 3 (extraction) reads all "
        "text in each tile verbatim; and Pass 4 (merge) is a text-only LLM call — with no image — "
        "that synthesises the observations into a caption. Because the model never sees the image "
        "in the final generation step, it physically cannot invent visual detail. This reduced "
        "hallucination to 0% and lifted extraction to 75% (see Section 6.2).")
    add_caption(doc, "Figure 12: Vision Multi-Pass Anti-Hallucination Pipeline", bookmark="fig_12")
    add_image(doc, "fig12_vision_multipass.png",
              "Figure 12: The four passes — structure, tile, extract, and an image-free merge — "
              "that let a compliant Gemma 3 4B read screenshots without fabricating detail.",
              width_in=6.2)

    add_heading(doc, "5.6  Aggregation, Alerts, and Notifications", level=2, bookmark="sec_5_6")
    add_para(doc,
        "Analyses are rolled up hourly and daily, per subreddit and per aspect, into an "
        "aggregates table that backs the dashboard KPIs. An alert engine detects volume spikes "
        "(> 2σ) and sentiment crashes (drop > 0.3 in 6 h) and classifies each triggering post "
        "into P1/P2. Notifications are routed by configurable groups: each group owns a set of "
        "subreddits and has its own email distribution list, Slack webhook, and priority filter, "
        "so different teams receive only the alerts they own. Every dispatch is written to an "
        "audit log, and a dry-run mode allows testing without sending.")
    add_caption(doc, "Figure 13: Alert Engine and Group-Based Notification Routing", bookmark="fig_13")
    add_image(doc, "fig13_alert_routing.png",
              "Figure 13: Aggregated analyses feed the alert engine, which classifies P1/P2, "
              "pushes to the live Alert Feed, and routes to subreddit-owned notification groups "
              "(email / Slack) with an audit log.", width_in=6.2)
    add_para(doc,
        "Alert Feed.", bold=True, space_after=2)
    add_para(doc,
        "The dashboard surfaces these alerts in a live Alert Feed. When the alert engine fires, "
        "the triggering event is pushed to the browser over the WebSocket channel (/ws/alerts) "
        "and prepended to the feed in real time, so the analyst sees a spike or sentiment crash "
        "the moment it is detected rather than on the next page refresh. Each feed entry shows the "
        "alert type (volume spike / sentiment crash / P1 negative), the affected subreddit and "
        "aspect, the metric that breached its threshold, a timestamp, and a one-click jump to the "
        "underlying posts in the Post Explorer. The feed is the operational front door: from a "
        "single alert an analyst can open the offending posts, send them to Review & Validate, "
        "and place them on the lifecycle board without leaving the workflow. WebSocket delivery "
        "with a polling fallback satisfies the real-time-alerts requirement (R8.1).")

    add_heading(doc, "5.7  Dashboard Overview", level=2, bookmark="sec_5_7")
    add_para(doc,
        "The React/FastAPI dashboard exposes eight surfaces: Brand Health (KPI tiles + trends), "
        "the real-time Alert Feed (Section 5.6), Post Explorer (full-text search and "
        "multi-filter), Review & Validate (the human-in-the-loop queue), Post Lifecycle (a Kanban "
        "board from triage to resolved), Insights & Competitor analysis, Pipeline Control, and "
        "Notifications. Real-time alerts are pushed over WebSocket, so the UI updates without "
        "polling. Figure 4 shows the information architecture and the correction feedback loop. "
        "Sections 5.8–5.13 describe the analyst-facing capabilities that were built after the "
        "mid-semester review and that complete the project.")
    add_caption(doc, "Figure 4: Dashboard Information Architecture and Feedback Loop", bookmark="fig_4")
    add_image(doc, "fig4_dashboard_map.png",
              "Figure 4: Dashboard pages, their backing tables/endpoints, and the correction "
              "feedback loop from Review & Validate into few-shot reply generation.")

    add_heading(doc, "5.8  Human-in-the-Loop Review & Validate", level=2, bookmark="sec_5_8")
    add_para(doc,
        "The Review & Validate page is the core operational workflow of the system and the point "
        "at which a human closes the loop on AI mistakes. It exists for two reasons: to guarantee "
        "that no incorrect or reputationally risky reply is ever sent to a customer, and — "
        "equally important — to capture, on every interaction, a small piece of high-quality "
        "supervised data that is stored for the future training of a domain model (Sections 5.9 "
        "and 5.10). In other words, the analyst is not only correcting the system; every "
        "correction is an investment that gradually reduces how much correcting the system will "
        "need in the future.")
    add_caption(doc, "Figure 5: Human-in-the-Loop Review & Validate Flow", bookmark="fig_5")
    add_image(doc, "fig5_review_validate_flow.png",
              "Figure 5: The Review & Validate flow — from the needs_review queue through "
              "correction, draft generation, posting or closing, to the feedback table that "
              "drives the learning loop.", width_in=5.2)
    add_para(doc, "Queue construction and ordering.", bold=True, space_after=2)
    add_para(doc,
        "The pending queue is built from the analyses table by selecting every post whose "
        "needs_review flag is set, and is ordered by priority so that the P1 tier (trust ≥ 0.70 "
        "and confidence ≥ 0.80) is surfaced before P2 (trust ≥ 0.50 and confidence ≥ 0.60). This "
        "means the analyst always works the highest-signal negative posts first. The queue is "
        "paginated with a 'Load More' control so that a large backlog does not slow the page, and "
        "a separate 'Reviewed' tab lets the analyst see what has already been handled. Enforcing "
        "needs_review at query time (rather than only when the queue is first loaded) is what "
        "guarantees that a post which has been confirmed or closed never reappears — a "
        "persistence defect that was diagnosed and fixed during post-midsem hardening.")
    add_para(doc, "Per-post correction controls.", bold=True, space_after=2)
    add_para(doc,
        "For each card the analyst can perform three independent corrections. (a) Sentiment: "
        "re-label the post as positive, neutral, or negative with a single click; the new label "
        "is written back to the analyses row. (b) Aspects: add or remove aspect tags, drawn from "
        "two separate sub-taxonomies — a customer set (pricing, product quality, customer support, "
        "store experience, online/app, delivery/pickup, returns) and an employee set "
        "(workforce/HR: pay, benefits, management, workload, safety) — so that, for example, an "
        "associate's pay complaint is never mislabelled as a customer pricing issue. (c) Trust: "
        "override the computed trust score when the analyst's judgement differs from the "
        "heuristic. Every one of these edits is persisted, and — critically — the (post, "
        "corrected-label, corrected-aspects) triple is written to the feedback table as a "
        "labelled example, building a growing, human-verified dataset that can later re-train or "
        "re-calibrate the sentiment and aspect models.")
    add_para(doc, "Reply and close flow.", bold=True, space_after=2)
    add_para(doc,
        "Once a post is understood, the analyst generates reply drafts (Section 5.9), edits the "
        "chosen draft, and clicks 'Post Reply'. Replies are never auto-posted: the action saves "
        "the reply to the audit log and, in live mode, posts it to Reddit through an OAuth surface "
        "gated by a dry-run flag. As soon as a reply is posted the card leaves the pending queue "
        "automatically and its record — the original post, the final approved reply, the model "
        "that produced the accepted draft, and any internal action note — is committed to the "
        "feedback table. A post can also be closed without a public reply along three explicit "
        "paths — 'reply sent / monitoring', 'action needed' (routed with an internal note to the "
        "actionable-items list), or 'no reply required' — and confirming any of these sets "
        "needs_review=false atomically so the queue stays clean.")

    add_heading(doc, "5.9  Smart Reply Composer — Dual-Draft (Rule + LLM)", level=2, bookmark="sec_5_9")
    add_para(doc,
        "For every negative post the Smart Reply Composer produces two candidate replies "
        "side-by-side so the analyst can compare and pick the better one, and — when the "
        "'Internal Action Note' checkbox (on by default) is enabled — a third, internal action "
        "recommendation for the operations team. The two customer-facing drafts come from "
        "deliberately different sources so that their strengths are complementary.")
    add_caption(doc, "Figure 6: Smart Reply Composer — Dual-Draft with Few-Shot Prompting", bookmark="fig_6")
    add_image(doc, "fig6_reply_cascade.png",
              "Figure 6: Two drafts are produced in parallel — a deterministic rule composer "
              "(Draft A) and a three-tier LLM cascade (Draft B: GPT-4o → Mistral → template) "
              "prompted with few-shot exemplars pulled from the feedback table.")
    add_para(doc, "Draft A — deterministic rule-based composer.", bold=True, space_after=2)
    add_para(doc,
        "The first draft is produced by a content-aware 'smart composer' that extracts the key "
        "entities and complaint keywords from the post (order type, product, error text from any "
        "vision caption, affected aspect) and assembles a reply from curated phrase pools with a "
        "randomised seed, so it varies on every call. It requires no model, runs instantly, and "
        "is always available — it is the guaranteed floor of quality even when every network and "
        "GPU path is down.")
    add_para(doc, "Draft B — LLM cascade with graceful degradation.", bold=True, space_after=2)
    add_para(doc,
        "The second draft is produced by a three-tier LLM cascade, each tier tried in order until "
        "one succeeds: (1) the Walmart LLM Gateway serving GPT-4o, called over HTTPS with the "
        "routing headers WM_CONSUMER.ID, WM_SVC.NAME, and WM_SVC.ENV; (2) if the gateway is "
        "unreachable, a local Mistral 7B Instruct model served by Ollama (temperature 0.55, "
        "num_predict 220); and (3) if neither LLM is available, a second pass of the smart "
        "composer with a different seed. The same cascade generates the internal action note "
        "(with a longer token budget and a lower temperature, since it is an operational "
        "instruction rather than a customer message). This design means the feature never fails "
        "closed: an analyst working fully offline still receives two usable drafts.")
    add_para(doc, "Few-shot prompting from approved history.", bold=True, space_after=2)
    add_para(doc,
        "The LLM tiers are not prompted cold. Before each generation the composer pulls the most "
        "recent analyst-approved replies (and their original posts) from the feedback table and "
        "injects them into the prompt as few-shot exemplars, together with the post's aspects and "
        "subreddit. The model is therefore shown 'here is how this brand has answered similar "
        "complaints before, now answer this one'. The practical effect is that drafts drift "
        "steadily towards the organisation's approved tone and policy with no model training at "
        "all — the analyst teaches the system simply by doing the normal job of approving "
        "replies. Because the exemplars are the actual approved replies, the improvement compounds "
        "as the feedback table grows.")

    add_heading(doc, "5.10  Learning Loop — Feedback for Future Retraining", level=2, bookmark="sec_5_10")
    add_para(doc,
        "The learning loop is the mechanism by which the system is designed to reduce its own "
        "dependence on the human reviewer over time. Every analyst action leaves a durable, "
        "structured record, and those records are consumed at two different time-scales.")
    add_caption(doc, "Figure 7: Feedback Learning Loop and Path to Auto-Reply", bookmark="fig_7")
    add_image(doc, "fig7_learning_loop.png",
              "Figure 7: Each approved reply enriches the feedback table, which feeds near-term "
              "few-shot prompting and ModernBERT re-calibration now, and a long-term Mistral "
              "fine-tune toward supervised / automatic reply once the corpus reaches ~1M pairs.")
    add_para(doc, "What is captured.", bold=True, space_after=2)
    add_para(doc,
        "On each reviewed post the feedback table stores: the original post text and any vision "
        "caption; the human-corrected sentiment and aspects; the trust override if any; the final "
        "approved reply text; which model (GPT-4o, Mistral, or the rule composer) produced the "
        "accepted draft; and the internal action note. Nothing is discarded — even a 'no reply "
        "required' close is a labelled negative example, and every action note is retained as a "
        "seed for future root-cause-analysis and auto-response capabilities.")
    add_para(doc, "Near-term use — zero-training few-shot improvement.", bold=True, space_after=2)
    add_para(doc,
        "In the immediate term this corpus feeds the few-shot prompting described in Section 5.9, "
        "so reply quality improves continuously without any training run. The same corrected "
        "labels can also be folded back into the ModernBERT fine-tuning set to periodically "
        "re-calibrate the sentiment classifier on genuinely new, human-verified examples.")
    add_para(doc, "Long-term use — supervised fine-tuning toward auto-reply.", bold=True, space_after=2)
    add_para(doc,
        "In the long term the feedback table is a supervised training set of (post → approved "
        "reply) pairs. Once it accumulates on the order of one million such pairs — a scale that "
        "becomes realistic as ingestion runs continuously across the 25 communities — the local "
        "Mistral-class model can be fine-tuned directly on this Walmart-specific data to generate "
        "high-quality replies on its own. At that point the workflow can shift from "
        "'draft → human approves → post' toward a supervised or fully automatic reply mode, with "
        "the human review queue reserved only for low-confidence or high-severity cases. This is "
        "the explicit end-state of the design: the human-in-the-loop stage exists both to keep "
        "today's replies safe and to manufacture the training data that will let a future model "
        "take over most of the routine work, progressively removing the HITL dependency rather "
        "than entrenching it.")

    add_heading(doc, "5.11  Post Explorer and Multi-Facet Search", level=2, bookmark="sec_5_11")
    add_para(doc,
        "The Post Explorer lets an analyst browse every analysed post with fast, multi-facet "
        "filtering. The available facets are sentiment (positive / neutral / negative), "
        "confidence and trust thresholds (slider controls), subreddit (multi-select over the 25 "
        "tracked communities), aspect, and date range (today / week / month / custom), combined "
        "with a free-text search over titles and bodies. Filters compose, so a query such as "
        "'negative + trust ≥ 0.6 + aspect = delivery_pickup + r/Sparkdriver + last 7 days' is a "
        "few clicks away. Each result card shows the title and a body excerpt, a colour-coded "
        "sentiment badge with confidence, the trust indicator, the aspect tags, the source "
        "subreddit and post time, a deep link that opens the original Reddit thread, and quick "
        "actions to send the post to Review & Validate or add it to the lifecycle board. The KPI "
        "tiles on the Brand Health page deep-link into the Explorer pre-filtered, so a spike in "
        "negative pricing posts is one click away from the underlying evidence, and the Alert Feed "
        "links here too — making the Explorer the common evidence surface for the whole dashboard.")

    add_heading(doc, "5.12  Post Lifecycle (Kanban Workflow)", level=2, bookmark="sec_5_12")
    add_para(doc,
        "The Post Lifecycle page is a Kanban board that tracks a complaint from first sighting to "
        "resolution across four states — Triaged → Acknowledged → In Progress → Resolved. New "
        "P1/P2 posts land in Triaged automatically; the analyst then acknowledges the case, starts "
        "work, and resolves it, dragging the card across the board. Every transition is written "
        "to the post_lifecycle table with a timestamp, which makes it possible to measure "
        "operational SLAs such as time-to-acknowledge and time-to-resolve, and to compute a "
        "resolution rate per aspect or per subreddit.")
    add_caption(doc, "Figure 8: Post Lifecycle Kanban Workflow", bookmark="fig_8")
    add_image(doc, "fig8_lifecycle_kanban.png",
              "Figure 8: The Kanban lifecycle — Triaged → Acknowledged → In Progress → Resolved "
              "with timestamped transitions, a two-step resolve, three close paths, and a "
              "three-day follow-up banner.", width_in=6.0)
    add_para(doc, "Two-step resolve and close paths.", bold=True, space_after=2)
    add_para(doc,
        "Resolution is a deliberate two-step flow so that the dashboard state can never claim a "
        "reply was posted before it actually was. In step one the analyst saves an action note "
        "and an optional drafted reply; in step two, after posting the reply on Reddit, the "
        "analyst marks the card Resolved. A card can be closed along three explicit paths that "
        "mirror the Review & Validate close options: 'reply sent' (acknowledged and now "
        "monitoring), 'action needed' (routed to the actionable-items list carrying its internal "
        "note), or 'no reply required'. A follow-up banner automatically flags any reply that has "
        "been sent three or more days ago without a recorded resolution, so cases that stall "
        "surface on their own instead of being forgotten. As with Review & Validate, each "
        "resolution enriches the feedback record that feeds the learning loop.")

    add_heading(doc, "5.13  Insights and Competitor Analysis", level=2, bookmark="sec_5_13")
    add_para(doc,
        "The Insights page turns the raw stream into strategic intelligence. It ranks the top "
        "negative issues by a volume × severity × recency score, grouped by aspect, with trend "
        "arrows showing whether each issue is improving or deteriorating week over week. An LLM "
        "summariser produces natural-language summaries of the week's dominant themes and "
        "suggested action items, and an emerging-topic detector surfaces new phrase clusters that "
        "appear across several posts in a short window — an early-warning signal for a problem "
        "that is beginning to trend. Together with a per-aspect drilldown (sentiment trend, "
        "volume, and representative posts), this gives product teams an aspect-level, "
        "evidence-backed picture rather than a raw feed.")
    add_para(doc, "Competitor Insights — cross-brand comparison.", bold=True, space_after=2)
    add_para(doc,
        "A dedicated competitor-pulse view benchmarks Walmart against Costco, Target, and Amazon "
        "on the same eight-aspect taxonomy, so a weakness is always expressed relative to peers "
        "rather than in isolation. It ingests the competitor subreddits (Appendix A) and scores "
        "them through the identical trust → sentiment → aspect pipeline, then presents: (i) a "
        "side-by-side sentiment index per brand and per aspect; (ii) cross-mentioned posts, where "
        "a single Reddit thread compares Walmart with a competitor (e.g. 'Walmart vs Costco "
        "pricing') — the highest-signal comparisons of all; (iii) per-subreddit breakdowns so an "
        "analyst can tell whether a gap is broad or concentrated in one community; and (iv) a "
        "share-of-negative-voice metric showing which brand is absorbing the most complaints in a "
        "given aspect and week. Because every brand is scored by exactly the same models, the "
        "comparison is apples-to-apples and the deltas are defensible.")

    add_heading(doc, "5.14  Storage Layer", level=2, bookmark="sec_5_14")
    add_para(doc,
        "Seven tables back the system: raw_posts (privacy-safe ingested data, partitioned by "
        "subreddit), analyses (sentiment, confidence, aspects, trust), aggregates (pre-computed "
        "KPIs by time window), alerts, feedback (human corrections, approved replies, and the "
        "model that produced each accepted draft — the corpus for the learning loop of Section "
        "5.10), notification_log, and post_lifecycle (the Kanban state and its full transition "
        "history). Reddit usernames are SHA-hashed before storage and a one-year retention "
        "default applies. SQLite is used in development; the same schema and partition keys "
        "(/subreddit for posts and analyses, /time_window for aggregates) map directly to Azure "
        "Cosmos DB for production, so moving to the cloud backend is a configuration swap rather "
        "than a rewrite.")

    # ══════════════ 6. EVALUATION ══════════════
    add_chapter(doc, "Evaluation and Results", bookmark="sec_6")

    add_heading(doc, "6.1  Sentiment Model Evaluation", level=2, bookmark="sec_6_1")
    add_para(doc,
        "The evaluation set is 200 hand-labelled long-form Walmart-Reddit posts (body 300–3,604 "
        "characters; class distribution neg=127, neu=65, pos=8). All numbers below are 5-fold "
        "out-of-fold cross-validation, so no sample is scored by a model that trained on it. "
        "Table 2 compares the fine-tuned ModernBERT against the RoBERTa baseline.")
    add_caption(doc, "Figure 14: Sentiment Macro-F1 — RoBERTa vs Fine-Tuned ModernBERT", bookmark="fig_14")
    add_image(doc, "fig14_sentiment_results.png",
              "Figure 14: Out-of-fold macro-F1 by length bucket — the fine-tuned ModernBERT "
              "recovers all seven ≥512-token posts (RoBERTa 5/7, ModernBERT 7/7) where RoBERTa's "
              "512-token cap truncates the decisive detail.", width_in=5.6)
    add_caption(doc, "Table 2: Sentiment Model Comparison (out-of-fold CV)", bookmark="tbl_2")
    add_table(doc,
        ["Model", "Context", "Macro-F1", "Notes"],
        [
            ["RoBERTa (twitter-sentiment-latest)", "512 tok", "0.6272", "Strong public baseline, tweet-trained"],
            ["ModernBERT-base (fine-tuned)", "1024 tok", "0.7642", "Production model, +0.137 macro-F1"],
            ["  Stage-3 5-fold CV mean", "1024 tok", "0.7362 ± 0.1155", "Domain specialization stage"],
        ],
        col_widths=[2.6, 0.9, 1.1, 2.1])
    add_para(doc,
        "The decisive gain is on long posts. Table 3 splits the evaluation at the RoBERTa "
        "512-token truncation boundary. All seven long posts are labelled negative by the human "
        "annotator, so the meaningful score for that bucket is a per-post correctness count "
        "rather than a per-class F1: ModernBERT recovers the two long posts the baseline "
        "mis-classifies (5/7 → 7/7 correct), showing the truncation ceiling that limited "
        "RoBERTa is removed.")
    add_caption(doc, "Table 3: Per-Length-Bucket Sentiment Results (5-fold OOF predictions)", bookmark="tbl_3")
    add_table(doc,
        ["Length bucket", "Baseline correct", "ModernBERT correct", "Recovered"],
        [
            ["< 512 tokens  (n=193)",                     "138/193 (72%)", "159/193 (82%)", "+21"],
            ["≥ 512 tokens  (n=7, all negative)",         "5/7",           "7/7",           "+2"],
            ["Overall  (n=200)",                           "143/200 (72%)", "166/200 (83%)", "+23"],
        ],
        col_widths=[2.6, 1.4, 1.4, 0.9])
    add_para(doc,
        "Honesty caveats, stated plainly: the labelled set is small (200) and AI-assisted with "
        "full human review (all suggestions accepted), the positive class is severely "
        "under-represented (n=8), and the 5-fold variance (±0.1155) is wide because each fold has "
        "few positives. The target 'stretch gate' of macro-F1 ≥ 0.80 was not fully reached "
        "(0.7642); this is reported transparently rather than by evaluating on training data, "
        "where the model trivially scores ~1.0. A blind re-check on fresh posts and a 3-seed "
        "ensemble are recommended (Section 10) to tighten and confirm these numbers.")

    add_heading(doc, "6.2  Vision Module Evaluation", level=2, bookmark="sec_6_2")
    add_para(doc,
        "The multi-pass pipeline was validated on 8 images initially and 25 images at scale. "
        "Table 4 summarises the before/after on the 8-image set.")
    add_caption(doc, "Table 4: Vision Module — Before vs After Multi-Pass", bookmark="tbl_4")
    add_table(doc,
        ["Metric", "Single-Pass", "Multi-Pass", "Change"],
        [
            ["Hallucination rate", "50% (4/8)", "0% (0/8)", "↓ 100%"],
            ["Overall failure rate", "75% (6/8)", "25% (2/8)", "↓ 67%"],
            ["Correct text extraction", "25% (2/8)", "75% (6/8)", "3× better"],
            ["Fabricated claims", "8", "0", "eliminated"],
            ["Avg latency / image", "~5 s", "~15 s", "3× (acceptable)"],
        ],
        col_widths=[2.2, 1.4, 1.4, 1.2])
    add_para(doc,
        "At scale (25 images), 22/25 passed cleanly, 3/25 were correct-but-sparse, and 0/25 "
        "failed, versus a 44% single-pass hallucination rate on the same set. Since ~80% of "
        "Walmart Reddit complaint images are screenshots or app screens — the category where "
        "single-pass fails hardest — this mitigation directly protects the integrity of the "
        "aggregated dashboard signal.")

    add_heading(doc, "6.3  Trust-Score Behaviour", level=2, bookmark="sec_6_3")
    add_para(doc,
        "On the 200-post file, mean trust was 0.712 and, because the fine-tuned model's "
        "confidence is near 1.0 on its own training data, the admission gate reduces to "
        "trust ≥ 0.30 and 100% of posts pass. This is expected training-fit behaviour; on unseen "
        "production posts confidence drops and the confidence factor begins to do real work. A "
        "sensitivity sweep of the five weightings (metadata-heavy, llm-heavy, equal, dedup-off) "
        "left the ranking of posts essentially unchanged, confirming that the exact weights are "
        "not a fragile optimum. An important data caveat: the free Arctic Shift provider returns "
        "account_age = 0 and karma = 0 for all posts, so on this file only text length and post "
        "score drive the metadata sub-score; age and karma will populate on the paid tier.")

    # ══════════════ 7. TOOLS ══════════════
    add_chapter(doc, "Tools, Technologies, and Configuration", bookmark="sec_7")
    add_caption(doc, "Table 5: Tools and Technologies", bookmark="tbl_5")
    add_table(doc,
        ["Layer", "Technology", "Key decision"],
        [
            ["Backend", "Python 3.13 + FastAPI + SQLite/Cosmos", "Free, local-first, modular"],
            ["Frontend", "React 18 + TypeScript + Vite + Tailwind", "Modern responsive SPA"],
            ["Sentiment", "ModernBERT (fine-tuned, 1024 tok)", "Domain-specialized, offline"],
            ["Aspects", "BART-MNLI / DeBERTa-v3 zero-shot", "No training data needed"],
            ["Vision", "Gemma 3 4B via Ollama (multi-pass)", "Policy-compliant, no hallucination"],
            ["Reply gen", "GPT-4o gateway → Mistral 7B → template", "Cascade fallback, HITL only"],
            ["Trust", "Metadata + dedup + LLM (weighted)", "Interpretable; flag, don't drop"],
            ["Ingestion", "Arctic Shift API (+ optional PRAW)", "Free, cursor-based, incremental"],
            ["Scheduling", "asyncio lifespan (6 h) + manual", "Cursor-based incremental sweep"],
            ["Observability", "structlog + cost ledger (JSONL)", "Per-call LLM cost tracking"],
        ],
        col_widths=[1.2, 2.6, 2.5])

    # ══════════════ 8. PROBLEMS ══════════════
    add_chapter(doc, "Problems Encountered and Mitigations", bookmark="sec_8")
    add_caption(doc, "Table 6: Problems Encountered and Mitigations", bookmark="tbl_6")
    add_table(doc,
        ["Problem", "Mitigation"],
        [
            ["Vision model hallucinated on 50% of images (fake receipts/prices)",
             "Four-pass pipeline with image-free final synthesis → hallucination 0%"],
            ["512-token models truncated long complaints, losing decisive detail",
             "Fine-tuned ModernBERT with 1024-token context → all seven ≥512-token posts recovered (RoBERTa 5/7, ModernBERT 7/7)"],
            ["Every state-of-the-art vision model was blocked by vendor policy",
             "Re-implemented paper techniques as a calling strategy on compliant Gemma 3 4B"],
            ["Free ingest provider returns age=0, karma=0 for all posts",
             "Metadata sub-score falls back to length/engagement; paid tier will populate"],
            ["Only 200 labelled posts; severe positive-class imbalance (n=8)",
             "Class weighting + oversampling; honest OOF CV; blind re-check planned"],
            ["Reviewed posts reappeared in the pending queue (persistence bug)",
             "Queue now always enforces needs_review=1; confirm/close set state atomically"],
            ["All replied posts routed to Actionable Items (auto action note)",
             "Made action-note generation opt-in; post-reply auto-closes as reply_sent"],
        ],
        col_widths=[3.0, 3.3])

    # ══════════════ 9. CONCLUSIONS ══════════════
    add_chapter(doc, "Conclusions and Recommendations", bookmark="sec_9")
    add_para(doc,
        "This dissertation delivered a complete, working prototype that converts a noisy public "
        "Reddit stream into a structured, aspect-tagged, trust-weighted retail brand-health feed, "
        "running entirely on a local, offline-first, zero-API-cost stack. The central technical "
        "claims are supported by honest evaluation: fine-tuning ModernBERT raised sentiment "
        "macro-F1 from 0.6272 to 0.7642 (out-of-fold), with all seven ≥512-token posts recovered "
        "(RoBERTa 5/7, ModernBERT 7/7) — the long-context motivation of ModernBERT is what "
        "delivers that recovery; and the multi-pass vision pipeline "
        "eliminated hallucination (50% → 0%) while tripling correct text extraction. The trust "
        "score and trust × confidence gate provide an interpretable, auditable filter that flags "
        "rather than drops low-credibility content, and the human-in-the-loop dashboard closes "
        "the loop between AI output and analyst action.")
    add_para(doc,
        "The main recommendation for anyone extending this work is to invest first in labelled "
        "data: the single largest limitation is the 200-post evaluation set with only eight "
        "positive examples. Beyond that, migrating storage to Azure Cosmos DB and moving "
        "account-metadata ingestion to a paid Reddit tier would let the trust score exercise its "
        "full metadata signal, and a small ensemble would tighten the sentiment variance. The "
        "architecture was deliberately built so that each of these is a configuration or data "
        "change rather than a redesign.")

    # ══════════════ 10. FUTURE WORK ══════════════
    add_chapter(doc, "Future Work", bookmark="sec_10")
    add_bullet(doc, "Azure Cosmos DB migration for production-grade, partitioned storage.")
    add_bullet(doc, "Twitter/X integration as a second data source once the Reddit pipeline is stable.")
    add_bullet(doc, "3-seed ModernBERT ensemble (+0.01–0.03 F1, tighter variance) and a blind "
                    "25-post re-check for defensibility.")
    add_bullet(doc, "Gemma 3 12B upgrade for the remaining edge-case images.")
    add_bullet(doc, "Expand the labelled set (especially the positive class) and add aspect-level "
                    "ground truth for aspect-F1 evaluation.")
    add_bullet(doc, "Spanish-language support for bilingual retail communities.")
    add_bullet(doc, "Azure AD authentication for multi-user access and an automated "
                    "feedback-loop retraining pipeline.")
    add_bullet(doc, "Move account-metadata ingestion to a paid Reddit tier so age/karma populate "
                    "the trust score.")
    add_bullet(doc, "Progressive HITL removal: continue harvesting approved post–reply pairs and "
                    "action notes, and once the corpus reaches ~1 million pairs, fine-tune the local "
                    "Mistral model on this Walmart-specific data to enable supervised or automatic "
                    "reply generation, reserving the review queue for low-confidence cases.")

    # ══════════════ 11. GLOSSARY AND ABBREVIATIONS ══════════════
    add_heading(doc, "Glossary and Abbreviations", level=1, bookmark="sec_11")
    add_para(doc,
        "The following technical terms and acronyms are used in this report; the section in which "
        "each is first introduced is indicated where relevant.", size=10, space_after=6)
    add_table(doc,
        ["Abbreviation / Term", "Expansion / Meaning"],
        [
            ["ABSA", "Aspect-Based Sentiment Analysis"],
            ["API", "Application Programming Interface"],
            ["CV", "Cross-Validation"],
            ["DocVQA", "Document Visual Question Answering (benchmark)"],
            ["F1", "Harmonic mean of precision and recall"],
            ["HITL", "Human-in-the-Loop"],
            ["KPI", "Key Performance Indicator"],
            ["LLM", "Large Language Model"],
            ["MPS", "Metal Performance Shaders (Apple-silicon GPU backend)"],
            ["NLI", "Natural Language Inference"],
            ["NLP", "Natural Language Processing"],
            ["OCR", "Optical Character Recognition"],
            ["OGP", "Online Grocery Pickup"],
            ["OOF", "Out-Of-Fold (cross-validation prediction)"],
            ["P1 / P2", "Priority tiers 1 and 2 for negative posts"],
            ["PRAW", "Python Reddit API Wrapper"],
            ["RSI", "Retail Sentiment Intelligence (this system)"],
            ["SPA", "Single-Page Application"],
            ["WILP", "Work Integrated Learning Programmes"],
        ],
        col_widths=[1.4, 4.9])

    # ══════════════ 12. REFERENCES ══════════════
    add_heading(doc, "References", level=1, bookmark="sec_12")
    add_para(doc,
        "References are listed in order of first citation. The serial number of each reference "
        "corresponds to the bracketed number [n] used in the body of the report (principally in "
        "Section 3, Literature Survey, and Section 5).", size=10, space_after=6)
    refs = [
        "Warner, B. et al. (2024). ModernBERT: A Modern Bidirectional Encoder. answer.ai / LightOn.",
        "Barbieri, F. et al. (2020). TweetEval: Unified Benchmark and Comparative Evaluation for "
        "Tweet Classification. Findings of EMNLP 2020.",
        "Demszky, D. et al. (2020). GoEmotions: A Dataset of Fine-Grained Emotions. ACL 2020.",
        "Yin, W., Hay, J., Roth, D. (2019). Benchmarking Zero-shot Text Classification via NLI. "
        "EMNLP-IJCNLP 2019.",
        "He, P. et al. (2021). DeBERTa: Decoding-enhanced BERT with Disentangled Attention. ICLR 2021.",
        "Google DeepMind (2025). Gemma 3 Technical Report.",
        "Ye, J. et al. (2023). UReader: Universal OCR-free Visually-situated Language "
        "Understanding. Findings of EMNLP 2023.",
        "Liu, Y. et al. (2024). TextMonkey: An OCR-Free Large Multimodal Model for Document "
        "Understanding.",
        "Hu, A. et al. (2024). mPLUG-DocOwl 1.5: Unified Structure Learning for OCR-free Document "
        "Understanding.",
        "Chen, Z. et al. (2024). InternVL2 / How Far Are We to GPT-4V? Shanghai AI Laboratory.",
        "Bai, S. et al. (2025). Qwen2.5-VL Technical Report. Alibaba.",
        "Reimers, N., Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese "
        "BERT-Networks (MiniLM family). EMNLP-IJCNLP 2019.",
        "World Economic Forum (2024). Fake online reviews and their economic impact.",
        "DataReportal (2024). Digital 2024: Global Overview Report.",
        "Sprout Social (2024). The Sprout Social Index.",
        "Reddit, Inc. (2024). Form S-1 Registration Statement (weekly active uniques).",
    ]
    for i, r in enumerate(refs, 1):
        add_para(doc, f"[{i}]  {r}", size=10, space_after=4)

    # ══════════════ APPENDIX A ══════════════
    add_page_break(doc)
    add_heading(doc, "Appendix A — Curated Subreddit Coverage", level=1, bookmark="app_a")
    add_caption(doc, "Table 7: Curated Subreddit Coverage (snapshot)", bookmark="tbl_7")
    add_table(doc,
        ["Group", "Representative subreddits"],
        [
            ["Walmart core", "r/walmart, r/walmartogp"],
            ["Spark drivers", "r/Sparkdriver"],
            ["OGP / backroom", "r/OGPBackroom"],
            ["Employees / HR", "r/WalmartEmployees"],
            ["Sam's Club", "r/samsclub"],
            ["Pharmacy", "r/WalmartPharmacy (and related)"],
            ["International", "Country-specific Walmart communities"],
            ["Competitors", "r/Costco, r/Target, r/Kroger, r/amazon (for benchmarking)"],
        ],
        col_widths=[1.8, 4.5])
    add_para(doc,
        "The full 25-community list with per-source counts is in "
        "data/reddit_walmart_communities.csv. The evaluation set of 200 posts drew from "
        "walmart=70, samsclub=44, Sparkdriver=30, WalmartEmployees=28, OGPBackroom=19, "
        "walmartogp=9.")

    # ══════════════ APPENDIX B ══════════════
    add_heading(doc, "Appendix B — Reproduction Commands", level=1, bookmark="app_b")
    add_para(doc, "Source code repository:", bold=True, space_after=2)
    add_para(doc,
        "https://gecgithub01.walmart.com/v0s01jh/Retail_Sentiment_Intelligence", size=10)
    add_para(doc, "Clone and set up the environment:", bold=True, space_after=2)
    add_para(doc,
        "git clone https://gecgithub01.walmart.com/v0s01jh/Retail_Sentiment_Intelligence.git && "
        "cd Retail_Sentiment_Intelligence && conda create -n rsi python=3.11 -y && "
        "conda activate rsi && pip install -r requirements.txt", size=10)
    add_para(doc, "Fine-tune ModernBERT (offline):", bold=True, space_after=2)
    add_para(doc,
        "HF_HUB_OFFLINE=1 TRANSFORMERS_OFFLINE=1 HF_DATASETS_OFFLINE=1 python "
        "scripts/train_modernbert_sentiment.py --stages 1,2,3 --folds 5 --max-length 1024 "
        "--batch-size 8", size=10)
    add_para(doc, "Honest evaluation (produces eval_results.json):", bold=True, space_after=2)
    add_para(doc, "python scripts/eval_sentiment_models.py", size=10)
    add_para(doc, "Run the full stack (API :8001, dashboard :3003, scheduler):", bold=True, space_after=2)
    add_para(doc, "./start.sh    (stop with ./start.sh stop)", size=10)
    add_para(doc,
        "Reference documentation in the repository: ARCHITECTURE.md, PIPELINE_AND_TOOLS.md, "
        "DASHBOARD_DESIGN.md, REQUIREMENTS.md, IMPLEMENTATION_PLAN.md, docs/MODEL_COMPARISON.md, "
        "docs/MODERNBERT_JOURNEY.md, and evaluation/TRUST_SCORE_REPORT.md.")

    # ══════════════ CHECKLIST (last page) ══════════════
    add_page_break(doc)
    add_heading(doc, "Checklist of Items for the Final Report", level=1, bookmark="sec_checklist")
    add_para(doc,
        "This checklist is completed, verified, and signed by the student and attached as the "
        "last page of the report, as required by the WILP guidelines.", italic=True, size=10)
    add_table(doc,
        ["S. No.", "Item", "Yes / No"],
        [
            ["1", "Is the final report neatly formatted with all the elements required for a "
                  "technical report?", "Yes"],
            ["2", "Is the Cover page in proper format (Annexure A)?", "Yes"],
            ["3", "Is the Title page (inner cover page) in proper format (Appendix B)?", "Yes"],
            ["4", "Is the Certificate from the Supervisor in proper format and signed?", "Yes"],
            ["5", "Is the Abstract within one page with technical keywords specified?", "Yes"],
            ["6", "Is the title of the report appropriate, descriptive, and precise?", "Yes"],
            ["7", "Is the list of abbreviations / acronyms included?", "Yes"],
            ["8", "Does the report contain a summary of the literature survey?", "Yes"],
            ["9", "Are pages, figures (title at bottom), tables (title at top), captions, and "
                  "appendices numbered properly?", "Yes"],
            ["10", "Is the conclusion of the report based on discussion of the work?", "Yes"],
            ["11", "Are references given at the end and cited properly inside the text?", "Yes"],
            ["12", "Is the report format and content according to the guidelines (not a mere "
                   "printout of a presentation or user manual; no source code)?", "Yes"],
        ],
        col_widths=[0.7, 4.6, 1.0])
    add_para(doc, "", space_after=8)
    add_para(doc, "Declaration by Student", bold=True, size=11, space_after=2)
    add_para(doc,
        "I certify that I have properly verified all the items in this checklist and ensure that "
        "the report is in proper format as specified in the course handout.", size=10, space_after=18)
    chk = doc.add_table(rows=2, cols=2)
    chk.autofit = True
    chk.cell(0, 0).text = "Place: Bengaluru"
    chk.cell(0, 1).text = "Signature of the Student"
    chk.cell(1, 0).text = "Date:"
    chk.cell(1, 1).text = "Name: Vishal Singh\nID No.: 2020AA05641"
    for row in chk.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
